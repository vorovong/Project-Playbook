"""
IR-기획서.md → IR-기획서.docx 변환 스크립트
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def create_ir_document():
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 스타일 설정 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # 한글 폰트 설정
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 제목 스타일
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = '맑은 고딕'
        h_font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)  # 딥 네이비
        h_rPr = h_style.element.get_or_add_rPr()
        h_rFonts = h_rPr.find(qn('w:rFonts'))
        if h_rFonts is None:
            h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
            h_rPr.append(h_rFonts)
        else:
            h_rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(24)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(12)

    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(18)
    doc.styles['Heading 2'].paragraph_format.space_after = Pt(8)

    doc.styles['Heading 3'].font.size = Pt(12)
    doc.styles['Heading 3'].paragraph_format.space_before = Pt(12)
    doc.styles['Heading 3'].paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════
    # 표지
    # ═══════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(주)신도시')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('The First City OS Provider.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    set_run_font(run, '맑은 고딕')

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('슈파(SUPA) — 사장님의 슈퍼 파트너')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4E, 0xCD, 0xC4)  # 민트
    set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('경영 지원 + 비용 절감, AI로 한 번에')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_run_font(run, '맑은 고딕')

    for _ in range(4):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Seed Round · 2.5억원')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026년 3월')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 목차
    # ═══════════════════════════════════════
    p = doc.add_heading('목차', level=1)

    toc_items = [
        ('1.', 'Problem — 숫자로 보는 위기'),
        ('2.', 'Solution — 슈파(SUPA): 사장님의 슈퍼 파트너'),
        ('3.', 'Product — 서비스 구조'),
        ('4.', 'Business Model — 수익 구조'),
        ('5.', 'Market — 시장 규모'),
        ('6.', 'Go-to-Market — 유저 확보 전략'),
        ('7.', 'Competitive Landscape — 경쟁 분석'),
        ('8.', 'Structure — PropCo / OpCo 구조'),
        ('9.', 'Team — AI 네이티브 조직'),
        ('10.', 'Traction — 현재까지'),
        ('11.', 'Ask — 투자 요청'),
        ('12.', '비전'),
        ('13.', 'Expected Synergy — 기대 효과'),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{num}  {title}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §2. Problem
    # ═══════════════════════════════════════
    doc.add_heading('1. Problem — 숫자로 보는 위기', level=1)

    add_quote(doc, '대한민국에서 가장 거대하지만, 가장 비효율적인 시장.')

    doc.add_heading('The 95% Majority', level=2)

    add_table(doc, [
        ['지표', '수치'],
        ['국내 소상공인 사업체 수', '약 790만 7천 개 (전체의 95.2%)'],
        ['소상공인 종사자 수', '약 1,090만 명 (전체의 45.9%)'],
    ])

    add_quote(doc, '소상공인은 국가 경제의 모세혈관이자 고용의 핵심이나,\n가장 낙후된 경영 환경에 방치되어 있다.')

    doc.add_heading('구조적 문제', level=2)

    add_table(doc, [
        ['문제', '현실'],
        ['경영 비효율', '세무, 마케팅, 재고, 고객관리를 사장님이 직접 → 본업에 집중 못 함'],
        ['고비용 구조', '개별 점포 단위의 구매·마케팅·채용 = 구조적으로 \'고비용 저효율\''],
        ['기존 지원의 한계', '교육·컨설팅(간접 지원) → "배웠는데 할 시간이 없다"'],
        ['디지털 격차', 'AI 기술은 발전하지만, 소상공인 현장에 맞는 도구는 거의 없음'],
        ['생존 위기', '인건비 급등 + AI 가속화 속에서, 데이터 없는 소상공인은 \'소멸\' 위기'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §3. Solution
    # ═══════════════════════════════════════
    doc.add_heading('2. Solution — 슈파(SUPA): 사장님의 슈퍼 파트너', level=1)

    add_quote(doc, '"뭉치면 싸지고, 쌓이면 보인다."\nAI로 경영을 도와주고, 유저가 모이면 비용까지 낮춰준다.\n사장님이 내는 돈은 거의 없다. AI가 운영하니까.')

    doc.add_heading('한 문장 정의', level=2)
    p = doc.add_paragraph()
    run = p.add_run('슈파 = 소상공인의 AI 경영 파트너. 경영 효율화(거의 무료) + 공동구매·물류 비용 절감.')
    run.bold = True
    set_run_font(run, '맑은 고딕')

    doc.add_heading('(주)신도시의 정의', level=2)
    add_quote(doc, '도시의 상업 활동을 관장하는 운영체제(OS)를 공급하는 로컬 테크(Local Tech) 기업.\n핵심 솔루션 \'슈파\'로 파편화된 로컬 비즈니스를 디지털로 혁신한다.')

    doc.add_heading('왜 지금 가능한가', level=2)
    items = [
        'AI 운영 비용 ≒ 0 → 수수료를 극도로 낮게 받아도 사업이 돌아감',
        '기존 SaaS는 구독료가 수익 → 무료로 못 함',
        '슈파는 구독료가 아니라 데이터와 네트워크 효과가 수익 → 무료/소액이 전략',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §4. Product
    # ═══════════════════════════════════════
    doc.add_heading('3. Product — 서비스 구조', level=1)

    doc.add_heading('3-1. 플라이휠 (핵심 도식)', level=2)

    add_diagram_box(doc,
        'AI로 운영비 ≒ 0  →  수수료 최소화\n'
        '      ↓\n'
        '사장님 부담 없이 가입  →  유저 확보\n'
        '      ↓\n'
        '유저가 모이면  →  공동구매·물류 원가 절감\n'
        '      ↓\n'
        '경영 지원 + 실제 비용 절감  →  더 많은 사장님 유입\n'
        '      ↓\n'
        '데이터 축적  →  진짜 자산\n'
        '      ↓\n'
        '      🔄 계속 돈다'
    )

    doc.add_heading('3-2. Phase 1 — 행궁동·수원 장악 (~27년 2분기)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('"Traffic First, Monetize Later." — 선 트래픽 확보, 후 고도화')
    run.bold = True
    run.font.italic = True
    set_run_font(run, '맑은 고딕')

    # Cost Innovation
    doc.add_heading('① Cost Innovation — 비용의 파괴적 혁신', level=3)
    p = doc.add_paragraph('초기 유저 확보를 위한 핵심 가치(Killer Value) = \'즉각적인 비용 절감\'')
    for run in p.runs:
        set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['기능', '설명'],
        ['슈파 Buying (공동 구매)', '1,000개 가맹점의 구매력 통합 → 식자재·소모품·배달용기를 프랜차이즈 본사 공급가 수준으로 인하'],
        ['Shared Service (공유 인프라)', '세무 기장, 법률 자문, 방역, 렌탈 등 공통 고정비를 \'N분의 1\' 구조로 최소화'],
    ])

    items = [
        '유통 도메인 파트너사와 연결 (직접 물류 안 함 = Asset Light)',
        '5개소만 묶어도 유통 업체에게 매력적 → 50개소, 100개소면 협상력 ↑',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    # AI 경영 효율화
    doc.add_heading('② AI 경영 효율화 (거의 무료)', level=3)

    p = doc.add_paragraph()
    run = p.add_run('MVP 핵심 — 26년 9월 출시')
    run.bold = True
    set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['기능', '설명'],
        ['AI 경영 대시보드', '매출·비용 자동 정리, 한눈에 보는 가게 현황'],
        ['세무·장부 자동 정리', '거래 자동 분류, 부가세·종소세 신고 보조'],
        ['SNS 콘텐츠 자동 생성', '가게 사진 → 인스타·네이버 포스트 자동 생성'],
        ['구매몰 (공동구매)', '식자재·소모품 공동구매 주문·결제 시스템'],
        ['고객·예약 관리', '단골 관리, 예약 자동 응대, 고객 알림'],
        ['AI 상담 챗봇', '사장님 질문에 AI가 365일 답변 (세무, 노무, 마케팅 등)'],
    ])

    p = doc.add_paragraph()
    run = p.add_run('MVP 확장 — 26년 12월까지')
    run.bold = True
    set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['기능', '설명'],
        ['매출·원가 분석 리포트', '월별·주별 매출 트렌드, 원가율 알림, 손익 분석'],
        ['마케팅 기본', '리뷰 자동 응대, 키워드 분석, 프로모션 제안'],
        ['인력 매칭', '단기 아르바이트·전문 인력 연결'],
        ['메뉴·상품 분석', '인기/비인기 메뉴 분석, 가격 최적화 제안'],
    ])

    p = doc.add_paragraph()
    run = p.add_run('프리미엄 기능 — Phase 2 (27년~)')
    run.bold = True
    set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['기능', '설명'],
        ['Predictive AI', '날씨·행사·유동인구 결합 → 발주량·필요 인력 예측'],
        ['Peer Benchmarking', '내 가게를 동일 상권 상위 10%와 비교'],
        ['맞춤 자동화', '가게별 고유 문제 진단 → 커스텀 해결 (현장 지식 기반)'],
        ['마케팅 고급 분석', '상권 분석, 타겟 마케팅, 광고 효율 추적'],
        ['입지·상권 분석', '창업 전 입지 평가, 상권 트렌드 리포트'],
        ['거래처·계약 관리', '납품 단가 비교, 계약 갱신 알림, 거래처 평가'],
    ])

    # 개발 방식
    p = doc.add_paragraph()
    run = p.add_run('개발 방식: ')
    run.bold = True
    set_run_font(run, '맑은 고딕')
    run = p.add_run('전략적 파트너와 공동 개발')
    set_run_font(run, '맑은 고딕')

    items = [
        '창업팀 3명(기획·현장·영업) + 전략적 파트너의 개발 인력',
        '창업팀은 기획·UX 설계·현장 검증·데이터 수집에 집중',
        '개발팀은 설계·구현·운영에 집중',
        '공동 개발 구조의 근거: 전략적 파트너는 슈파의 소상공인 현장 채널과 운영 데이터에 접근하고, 신도시는 개발 역량을 확보하는 상호 이익 구조',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    # 완전 무료 전략
    doc.add_heading('③ 완전 무료 전략 + 프리미엄 쿠폰', level=3)

    add_quote(doc, 'Phase 1에서는 구독료 0원. 경영자동화 전 기능을 무료로 개방한다.\n프리미엄 기능도 홍보 쿠폰으로 무료 제공.\n→ 유저가 프리미엄에 익숙해진 상태로 Phase 2 진입 → 구독 전환율 극대화.')

    add_table(doc, [
        ['기존 플랫폼', '수수료', '슈파'],
        ['네이버 스마트스토어', '5~8%', '구매몰 수수료 3%'],
        ['쿠팡 마켓플레이스', '10~15%', '(기존 대비 50~80% 절감)'],
        ['배달의민족', '6.8% + 광고비 (실질 15~20%)', ''],
        ['SPC 등 B2B 유통', '15~25%', ''],
    ])

    items = [
        '슈파 구매몰 수수료 3% 중 PG 결제 수수료 약 1.5% → 슈파 실질 마진 약 1.5%',
        'AI 운영비 ≒ 0이므로 1.5% 마진으로도 운영 가능',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    run = p.add_run('KPI: ')
    run.bold = True
    set_run_font(run, '맑은 고딕')
    run = p.add_run('행궁동 100개소 → 수원 1,000개소')
    set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    run = p.add_run('수익: ')
    run.bold = True
    set_run_font(run, '맑은 고딕')
    run = p.add_run('구매몰 수수료 3% + B2G 용역 (캐시카우). 구독료 없음')
    set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # Phase 2
    doc.add_heading('3-3. Phase 2 — 국내 확장 + 수익화 (27년~)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('세 가지 수익 엔진이 동시에 켜지는 단계')
    run.bold = True
    set_run_font(run, '맑은 고딕')

    doc.add_heading('① 구독 서비스 시작 — 월 4,900원', level=3)

    p = doc.add_paragraph('Phase 1에서 프리미엄 쿠폰으로 무료 체험한 유저 → 자연스러운 유료 전환.')
    for run in p.runs:
        set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['구분', '내용'],
        ['무료 (유지)', '기본 경영자동화 (세무 자동정리, 매출분석, 기본 리포트, SNS 기본 생성)'],
        ['프리미엄 (월 4,900원)', 'Predictive AI, 맞춤 자동화, 마케팅 고급 분석, Peer Benchmarking, 우선 지원'],
        ['광고 (별도)', '유통사/공급사 → 소상공인 대상 타겟 광고'],
    ])

    add_quote(doc, '전환율 30% 근거\n\n'
        '• 가격: 월 4,900원 — 5,000원 미만 최소 저항 구간\n'
        '• 체험 선행: Phase 1에서 프리미엄을 이미 무료로 써봄\n'
        '• 가치 체감: Predictive AI = 직접적 매출 영향\n'
        '• 명분: 구독료의 일부가 상권 기금으로\n'
        '• 벤치마크: Spotify 무료→유료 44%, Slack 30%+')

    # 상권 기금
    doc.add_heading('② 상권 기금 선순환 — 구독이 동네를 지킨다', level=3)

    add_diagram_box(doc,
        '소상공인이 월 4,900원 구독\n'
        '      ↓\n'
        '일부를 공존 기금에 출연\n'
        '      ↓\n'
        '기금이 로컬 기획 프로젝트 지원\n'
        '(상권 활성화, 젠트리피케이션 방지, 재기 지원)\n'
        '      ↓\n'
        '상권이 좋아지면 → 소상공인 매출 ↑\n'
        '      ↓\n'
        '더 많은 소상공인이 슈파 가입 + 구독 유지\n'
        '      ↓\n'
        '      🔄 선순환'
    )

    add_quote(doc, '기존 플랫폼과의 차이: 배민·쿠팡은 수수료를 가져가기만 한다.\n슈파는 돈이 지역에서 도는 구조를 만드는 플랫폼이다.')

    # 데이터 판매
    doc.add_heading('③ 데이터 판매 — 진짜 자산이 되는 단계', level=3)

    items = [
        '1,000개소의 실데이터(Real Data)가 모여야 유의미한 AI 학습과 패턴 분석 가능',
        '전략적 파트너와 데이터 판매 → 실제 돈 버는 구조',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['데이터 구매자', '활용', '예상 가격대'],
        ['금융사', '소상공인 신용평가, 맞춤 대출', '연 3,000만~1억원'],
        ['유통사', '상권·소비 트렌드, 수요 예측', '연 2,000~8,000만원'],
        ['보험사', '업종별 리스크 분석', '연 1,000~3,000만원'],
        ['지자체', '상권 정책 근거 데이터', '연 500~2,000만원'],
        ['프랜차이즈 본사', '벤치마킹, 신규 출점', '리포트당 30~100만원'],
    ])

    add_quote(doc, '데이터 가치의 핵심: 매출 데이터는 카드사에서도 얻을 수 있지만, 원가·인건비·재고·소비자 행동 데이터는 사업장 내부에서만 나온다. 이 "내부 운영 데이터"가 프리미엄.')
    add_quote(doc, '레퍼런스: 캐시노트(KCD)는 170만 사업장 기반 연결 매출 1,428억원. 사업장당 약 8만원/년 수준.')

    doc.add_page_break()

    # Phase 3
    doc.add_heading('3-4. Phase 3 — 피지컬 AI (라스트마일)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('현실에서 데이터를 모으고, 피지컬 AI를 구현하는 도메인')
    run.bold = True
    set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph('Phase 1~2에서 쌓은 것:')
    for run in p.runs:
        set_run_font(run, '맑은 고딕')
    items = [
        '소상공인 현장 접점 (설치·운영 채널)',
        '업종별 운영 데이터 (학습 데이터)',
        '물류·배송 네트워크 (인프라)',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph('→ 피지컬 AI 접목: 배달·물류 로봇, 매장 자동화 기기, 무인 주문·결제, 재고 자동 관리')
    for run in p.runs:
        set_run_font(run, '맑은 고딕')

    add_quote(doc, '핵심: 로봇을 만드는 건 다른 회사도 하지만, 소상공인 현장에 놓을 수 있는 도메인(현장 채널 + 데이터 + 현장 지식)을 가진 곳은 우리뿐.')

    # Ecosystem Fund
    doc.add_heading('3-5. Ecosystem Fund (공존 기금)', level=2)

    add_quote(doc, '"내 구독료가 우리 동네를 지킨다"')

    p = doc.add_paragraph('슈파는 단순 도구가 아니라 지역 생태계 플랫폼이다. 플랫폼 수익의 일부를 공존 기금으로 적립하여, 소상공인이 만든 가치가 소상공인에게 돌아가는 선순환을 만든다.')
    for run in p.runs:
        set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    run = p.add_run('재원: ')
    run.bold = True
    set_run_font(run, '맑은 고딕')
    run = p.add_run('프리미엄 구독료의 일부 + 구매몰 수수료 마진의 일부')
    set_run_font(run, '맑은 고딕')

    add_table(doc, [
        ['영역', '예시'],
        ['상권 활성화', '로컬 기획 프로젝트, 공동 마케팅'],
        ['젠트리피케이션 방지', '임대료 안정 기금, 장기 임차 지원'],
        ['재기 지원', '폐업 소상공인 재창업 자금, 컨설팅'],
    ])

    p = doc.add_paragraph()
    run = p.add_run('운용 구조: ')
    run.bold = True
    set_run_font(run, '맑은 고딕')
    items = [
        '적립률·집행 기준은 Phase 2(전국 확장) 시작 시 확정',
        '지역 소상공인 대표 + 신도시 + 외부 자문으로 운영위원회 구성',
        '투명한 사용 내역 공개 (연 1회 이상)',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    add_quote(doc, '투자자 관점에서의 의미: 공존 기금이 만드는 \'지역 귀속감\'은 경쟁사가 복제할 수 없는 유저 잔존율의 핵심 동력이다. 비용이 아니라 해자(moat).')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §5. Business Model
    # ═══════════════════════════════════════
    doc.add_heading('4. Business Model — 수익 구조', level=1)

    doc.add_heading('Two-Track: Cash Cow + Tech Growth', level=2)

    add_table(doc, [
        ['Track', '수익원', '시점'],
        ['Cash Cow (생존)', 'B2G 용역·컨설팅 (글로컬상권 수행 경험 기반, 연 1~2억)', 'Y1~'],
        ['Cash Cow (생존)', '팝업/이벤트 대행 (공존공간 인프라 활용)', 'Y1~'],
        ['Tech Growth', '구매몰 수수료 3% (구독료 없음. 선 트래픽 확보)', 'Phase 1'],
        ['Tech Growth', '구독 월 4,900원 (전환율 30%) + 광고', 'Phase 2'],
        ['Tech Growth', '데이터 판매 (B2B)', 'Phase 2'],
        ['Tech Growth', '피지컬 AI (RaaS, 하드웨어)', 'Phase 3'],
    ])

    doc.add_heading('왜 무료/소액이 가능한가', level=2)

    add_diagram_box(doc,
        '기존 구독형 서비스                   슈파\n'
        '━━━━━━━━━━━━               ━━━━━━━━━━\n'
        '사람이 운영 → 인건비 큼           AI가 운영 → 운영비 ≒ 0\n'
        '구독료가 수익 → 무료 불가         데이터·네트워크가 수익 → 무료가 전략\n'
        '유저 이탈 = 매출 하락            유저 증가 = 데이터 자산 + 물류 협상력 ↑'
    )

    doc.add_heading('3개년 추정 손익 (P&L)', level=2)

    p = doc.add_paragraph('(단위: 억원)')
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_run_font(p.runs[0], '맑은 고딕')

    add_table(doc, [
        ['구분', '2026년', '2027년', '2028년'],
        ['매출', '2.1', '5.8', '27.4'],
        ['  - B2G 용역 (캐시카우)', '2.0', '2.0', '2.0'],
        ['  - 구매몰 수수료 3%', '0.1', '2.1', '18.9'],
        ['  - 구독 (월 4,900원)', '—', '0.1', '0.5'],
        ['  - 데이터 판매', '—', '0.8', '3.0'],
        ['  - 광고', '—', '0.9', '3.0'],
        ['비용', '3.0', '5.5', '11.0'],
        ['영업이익', '-0.9', '+0.3', '+16.4'],
    ], highlight_rows=[1, 7, 8])

    items = [
        '2026년 (기반 구축): B2G 용역으로 버티면서 유저 확보에 올인. 구독료 없음. 적자 9,000만원은 투자금으로 커버',
        '2027년 (전환점): 구매몰 수수료 + 데이터 판매 + 광고가 동시에 켜짐. 흑자 전환',
        '2028년 (폭발적 성장): 전국 5,000개소로 확장. 구매몰 거래액이 레버리지 → 영업이익 16억',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    add_quote(doc, '핵심 매출 엔진 = 구매몰 수수료 3%\n가게가 늘면 → 구매몰 거래가 늘고 → 수수료가 자동으로 커지는 구조.\n사람을 더 안 뽑아도 매출이 올라감 (= Asset Light의 힘).\n구독(월 4,900원)은 부가 수입이지, 핵심 수익이 아님.')

    add_quote(doc, '2028년 구매몰 수수료 18.9억원 산출\n\n'
        '전국 5,000개소 × 구매몰 이용률 35%(보수적)\n'
        '= 1,750개소 이용\n'
        '× 전 업종 가중 월 구매액 300만원 × 12개월\n'
        '= 연간 GMV 약 630억원\n'
        '× 수수료 3%\n'
        '= 18.9억원\n\n'
        '최저가 정책 효과로 이용률이 50%까지 상승하면 GMV 약 900억원, 수수료 약 27억원.')

    doc.add_heading('P&L 추정 근거', level=3)

    add_table(doc, [
        ['가정', '수치', '출처'],
        ['음식점 월 식재료비', '매출의 40~48% (월 760~910만원)', 'KREI 외식업체 경영실태조사 2024'],
        ['소상공인 월 평균 매출', '1,660만원', '중기부 소상공인실태조사 2023'],
        ['행궁동 식음 업종 비율\n(장안동·신풍동 기준)', '54% (449개 중 244개)', '수원시 상권컨설팅 2023'],
        ['전 업종 가중 월 구매액', '250~350만원', '식음(700만+) + 비식음(100만) 가중 추정'],
        ['구매몰 이용률', '35% → 50% (최저가 효과)', '자체 추정'],
        ['구독 전환율', '30%', 'Spotify 44%, Slack 30%+ 벤치마크'],
        ['프리미엄 가격', '월 4,900원', '네이버 플러스 멤버십 동일가 참조'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §6. Market
    # ═══════════════════════════════════════
    doc.add_heading('5. Market — 시장 규모', level=1)

    doc.add_heading('슈파의 시장은 SaaS가 아니라 유통이다', level=2)
    add_quote(doc, 'AI 경영보조는 유저를 모으는 무료 서비스이고, 진짜 매출은 구매몰 수수료에서 나온다.\n따라서 TAM은 SaaS 시장(4,000~6,000억원)이 아니라, 소상공인 구매/유통 시장이다.')

    doc.add_heading('TAM · SAM · SOM', level=2)

    add_table(doc, [
        ['구분', '정의', '규모'],
        ['TAM', '국내 소상공인 구매/유통 시장', '64조원'],
        ['', '→ 슈파 수수료 기준 (64조 × 3%)', '약 1.9조원'],
        ['SAM', '수원시 소상공인 약 11~13만개', '약 5,000~7,000억원'],
        ['SOM', '행궁동 100개소 → 수원 1,000개소', 'Phase 1 GMV 약 40~70억원'],
    ])

    doc.add_heading('TAM 산출 근거', level=3)

    add_table(doc, [
        ['항목', '수치', '출처'],
        ['국내 B2B 식자재 유통 시장', '64조원 (2025년 추정)', '벤처스퀘어, 한국경제'],
        ['국내 소상공인 수', '596만 개 (2023년 기준)', '중기부 소상공인실태조사 2023'],
        ['음식점 월 식재료비', '매출의 40~48%', 'KREI 외식업체 경영실태조사 2024'],
        ['수원시 소상공인 수', '약 11~13만 개', '경기도 사업체 수 × 수원 인구 비중'],
        ['행궁동 사업체 수\n(조사 구역: 장안동·신풍동)', '449개 (식음 54%)', '수원시 상권컨설팅 2023'],
        ['행궁동 상권 전체\n(구천동·남문 일대 포함)', '약 2,400개 사업자', '현장 파악'],
    ])

    doc.add_heading('시장 트렌드', level=2)
    items = [
        '소상공인 디지털 전환 가속화 — 정부 데이터바우처 사업 등 정책 지원 확대',
        'AI 비용 급락 → 소상공인도 AI를 쓸 수 있는 시대 도래',
        'B2B 식자재 유통 시장 디지털화 가속 — 80% 이상이 아직 영세 유통업체 중심',
        '피지컬 AI(배달로봇 등) 규제 완화 추세',
        '한국 데이터 시장: 27조원(2023) → 49조원(2028 목표), 연평균 12.7% 성장',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §7. GTM
    # ═══════════════════════════════════════
    doc.add_heading('6. Go-to-Market — 유저 확보 전략', level=1)

    doc.add_heading('"Zero Barrier." — 진입 장벽 제거', level=2)

    add_diagram_box(doc,
        'Phase 1-1: 행궁동 100개소                    (~26년 12월)\n'
        '━━━━━━━━━━━━━━━━━━━━━━\n'
        '채널 ① 글로컬상권 네트워크             → 초기 50개소+\n'
        '       (2년 142회 현장 협업 관계)\n'
        '       200여 개 핵심 사업자 DB 기보유\n'
        '채널 ② 찾아가는 온보딩                → 100개소 확대\n'
        '       공존공간 2층 HQ를 활용한 오프라인 지원\n'
        '\n'
        'Phase 1-2: 수원 1,000개소                    (~27년 2분기)\n'
        '━━━━━━━━━━━━━━━━━━━━━━\n'
        '채널 ③ 행정 공식 접근                  → 전통시장·상권 단위 확보\n'
        '       "무료 AI 경영 지원" = 지자체 정책과 일치\n'
        '채널 ④ 전략적 파트너 네트워크           → 유통 업체 연결\n'
        '\n'
        'Phase 2: 전국                               (27년 1분기~)\n'
        '━━━━━━━━━━━━━━━━━━━━━━\n'
        '채널 ⑤ 수원 성공 모델 → 타 지자체 복제\n'
        '채널 ⑥ 데이터 판매 실적 기반 파트너 확대'
    )

    doc.add_heading('핵심 논리', level=2)
    items = [
        '무료 서비스 = 행정 접근 가능. 지자체 입장에서 "소상공인에게 무료 AI 경영 지원"은 정책 목표와 일치',
        '공존공간 대표가 이미 2년간 지자체·소상공인·행정 3자 협업을 해본 사람',
        '민간 주도 상권관리기구(BID) 행궁동행의 기획 주체 → 상권 내 \'관리자 명분\' 보유',
        '다른 스타트업이 "행정에 접근하겠습니다"라고 하면 뜬소리지만, 이 팀은 이미 하고 있는 사람들',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §8. Competitive
    # ═══════════════════════════════════════
    doc.add_heading('7. Competitive Landscape — 경쟁 분석', level=1)

    doc.add_heading('포지셔닝 맵', level=2)

    add_diagram_box(doc,
        '                   AI 깊이 깊음\n'
        '                       ↑\n'
        '                       │\n'
        '                       │    ★ 슈파\n'
        '                       │    (AI 경영 자동화\n'
        '                       │     + 공동구매/물류\n'
        '                       │     + 데이터 + 피지컬 AI)\n'
        '                       │\n'
        '  외식업 한정 ←────────┼────────→ 전 업종\n'
        '                       │\n'
        '        스마트푸드      │    캐시노트\n'
        '        네트웍스        │    (매출관리\n'
        '        (식자재 유통    │     + 금융중개)\n'
        '         + 장부앱)      │\n'
        '                       ↓\n'
        '                   AI 깊이 얕음'
    )

    doc.add_heading('주요 비교', level=2)

    add_table(doc, [
        ['구분', '캐시노트 (KCD)', '스마트푸드네트웍스', '슈파 (SUPA)'],
        ['핵심', '무료 매출관리 → 금융중개', '식자재 유통 + 무료 장부앱', 'AI 경영 자동화 + 공동구매'],
        ['유저', '170만 사업장', '35,000개 매장', '0 (MVP 전)'],
        ['업종', '전 업종', '외식업만', '전 업종'],
        ['AI 깊이', '매출 분석', '리뷰 자동 응답 1개', '경영 전반 AI 자동화 + Predictive AI'],
        ['비용 절감', '금융 중개', '식자재 원가만', '경영비 + 물류 + 공동구매 + Shared Service'],
        ['물류', '❌', '자체 물류 (Asset Heavy)', '파트너 연결 (Asset Light)'],
        ['데이터 수익', '금융중개형', '❌', 'B2B 데이터 판매'],
        ['피지컬 AI', '❌', '❌', 'Phase 3 로드맵'],
    ])

    doc.add_heading('MOAT (경쟁 우위)', level=2)

    add_table(doc, [
        ['우위', '설명'],
        ['현장 지식', '2년 142회 현장 협업 + 200여 핵심 사업자 DB. 기술만으론 복제 불가'],
        ['AI 내재화 조직', '3명 + 슈파 AI. 운영비 최소화 = 낮은 수수료 = 유저 확보'],
        ['관리자 명분', '행궁동행(BID) 기획 주체. 진입 장벽 높은 로컬 시장에서 \'관리자\'로 시작'],
        ['행정 채널', '지자체와 2년 협업 + B2G 누적 16.4억원 수행 실적'],
        ['전략적 파트너십', '유통 도메인 연결 + 데이터 판매 협업'],
        ['플라이휠', '유저 → 공동구매 → 비용절감 → 유저. 먼저 돌린 사람이 이김'],
        ['Asset Light', '제조 리스크 공존공간에 잔존, 신도시는 재고 0·설비 0의 순수 Tech/Platform'],
    ])

    add_quote(doc, '검증된 선례: 스마트푸드네트웍스가 "무료 앱 → 유저 확보 → 물류 수익" 플라이휠로 35,000개 매장을 확보했다.\n슈파는 이 검증된 구조에 AI 경영 자동화 + 전 업종 확장 + 데이터 수익 + 피지컬 AI를 얹는다.')

    doc.add_heading('슈파는 SaaS가 아니다 — 글로벌 패러다임 전환', level=2)

    p = doc.add_paragraph('2026년, 세계 소프트웨어 시장에 "SaaSpocalypse"(SaaS 종말)라는 말이 나왔다. SaaS 기업 주식에서 시가총액 약 1,300조원이 증발했고, 투자자들은 "좌석당 구독 모델"에서 등을 돌리고 있다. (출처: TechCrunch, 2026.02)')
    for run in p.runs:
        set_run_font(run, '맑은 고딕')

    add_quote(doc, '"AI는 소프트웨어 시장(400조)이 아니라 노동 시장(1경 7,000조)을 먹는다." — a16z (2025)\n\n"우리는 더 이상 소프트웨어를 배포하는 것이 아니다. 지능을 배포하는 것이다." — NFX (2025)')

    doc.add_heading('글로벌 선례: 구독이 아니라 거래에서 돈 번 기업들', level=3)

    add_table(doc, [
        ['기업', '시작', '진화', '수익 비율'],
        ['Toast (미국, 33조원)', '외식업 POS 도구', '금융+데이터 플랫폼', '구독 1 : 금융 6'],
        ['Shopify (캐나다)', '온라인 스토어 빌더', '결제+AI 커머스 플랫폼', '구독 1 : 결제 3'],
        ['캐시노트 (한국)', '무료 매출관리', '금융 중개', '무료 → 연결 매출 1,428억'],
    ])

    add_quote(doc, '슈파는 SaaS가 아니다.\nAI 경영 서비스를 무료로 제공해서 유저를 모으고, 유저의 구매력을 통합해 거래 수수료로 수익을 만드는 AI 기반 커머스 플랫폼이다.\n구독 모델이 아니기 때문에, SaaSpocalypse에 면역인 구조다.')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §9. Structure
    # ═══════════════════════════════════════
    doc.add_heading('8. Structure — PropCo / OpCo 구조', level=1)

    doc.add_heading('분리 구조', level=2)

    add_diagram_box(doc,
        '(주)공존공간 [PropCo — Asset & Governance]\n'
        '├── 부동산 자산 보유, 임대\n'
        '├── 지역 거버넌스 및 파트너십 관리\n'
        '├── 제조업(양조) 리스크 잔존\n'
        '└── 신도시에 \'실증 단지(Test-bed)\' 접근권 제공\n'
        '\n'
        '(주)신도시 [OpCo — Tech & Operation]\n'
        '├── 슈파(SUPA) 개발·운영\n'
        '├── City OS 구축\n'
        '├── 지역창업스튜디오 (비즈니스 스케일업)\n'
        '└── Asset-Light: 재고 0, 설비 0, 오직 기술과 마케팅에 집중'
    )

    doc.add_heading('전략적 연합', level=2)

    add_table(doc, [
        ['교환', '공존공간 → 신도시', '신도시 → 공존공간'],
        ['데이터', '15년 축적 아날로그 상권 데이터 + 300여 파트너 DB', '실시간 디지털 상권 분석 데이터 (부동산 밸류업)'],
        ['인프라', '물리적 거점 (오피스, 라운지, 팝업 공간)', '정당한 임대료 + 인프라 사용료'],
        ['실증', '슈파 솔루션 테스트베드 우선 활용', '상권 매력도 유지·상승에 기여'],
    ])

    doc.add_heading('투자자 관점 이점', level=2)
    items = [
        '부동산 등 무거운 자산에 자금이 묶이지 않음',
        '오직 성장(Growth)과 기술(Tech)에 자금 투입',
        '높은 ROE(자기자본이익률) 달성 가능',
        'Clean Company: 부채와 리스크 제거된 상태에서 출발',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §10. Team
    # ═══════════════════════════════════════
    doc.add_heading('9. Team — AI 네이티브 조직', level=1)

    doc.add_heading('3명 + AI = 10인 이상의 속도', level=2)

    add_quote(doc, '이 팀의 경쟁력은 개인 이력이 아니라, AI를 내재화한 조직 구조다.\n기획·전략·영업·세무·마케팅·데이터 분석을 3명이 AI와 함께 처리한다.\n사람을 늘리는 게 아니라, AI로 속도를 만드는 조직.')

    add_table(doc, [
        ['', '박승현', '김민호', '이승훈', '🤖 슈파 AI'],
        ['역할', '대표 · 기획 총괄', 'AI 전략 · 시스템 설계', '운영 · 영업', 'AI 에이전트'],
        ['도메인', '문화기획 15년 · F&B 10년', '공존공간 코파운더', '무역업 · 소상공인 컨설턴트 10년', '사장님 경영 보조'],
        ['AI 활용', 'AI 기획 자동화 · IR/사업계획서 AI 협업', 'AI 업무 자동화 설계 · 프롬프트 엔지니어링', 'AI 기반 영업 데이터 분석 · 고객 관리 자동화', '세무·마케팅·분석 자동 처리'],
        ['핵심 실적', '글로컬상권 2년 142회 현장 협업\nB2G 누적 16.4억원 PM\n로컬 브랜드 운영\n🏆 2024 중기부장관 표창\n🏆 2025 로컬프랜드페어 AI 부문 표창', 'AI 활용 업무 프로세스 구축', '소상공인 현장 네트워크 구축', '—'],
    ])

    doc.add_heading('팀의 핵심 강점', level=2)

    add_quote(doc, '"AI 생태계에서 성장하는 조직"')

    add_table(doc, [
        ['강점', '설명'],
        ['AI 내재화', 'AI를 도구로 쓰는 게 아니라, AI가 팀원이다. 3명이 10명분의 속도를 낸다'],
        ['현장 × 기술', '15년 로컬 현장 지식 + AI 자동화 = 다른 테크 스타트업이 못 가진 조합'],
        ['검증된 실행력', 'B2G 16.4억원 완주 + 142회 현장 협업 + 중기부장관 표창 2회'],
        ['확장 가능한 구조', '외부 개발팀 공동 개발 + AI 자동화 → 인원 추가 없이 서비스 확장 가능'],
    ])

    add_quote(doc, '투자자에게 보이는 것: 이 팀은 돈을 태워서 사람을 늘리는 조직이 아니다.\nAI로 효율을 만들고, 현장 지식으로 진입 장벽을 세우고, 검증된 실행력으로 결과를 낸다.')

    doc.add_heading('지역창업스튜디오 (Company Builder)', level=2)
    items = [
        '단순 투자를 넘어, 성공 확률을 설계하는 컴퍼니 빌더',
        'City OS에 축적된 상권 데이터를 예비 창업자에게 제공',
        '공존공간 유휴 공간을 활용한 \'팝업 테스트\' 선행 → 본 창업 전 시장성 검증',
        '스튜디오를 통해 성장한 기업은 다시 슈파의 핵심 유저 → 플라이휠 강화',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, '맑은 고딕')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §11. Traction
    # ═══════════════════════════════════════
    doc.add_heading('10. Traction — 현재까지', level=1)

    add_table(doc, [
        ['항목', '상태'],
        ['법인 설립', '2026년 3월 (진행 중)'],
        ['슈파 MVP', '2026년 9월 목표'],
        ['소상공인 네트워크', '행궁동 상권 전체 약 2,400 사업자(구천동·남문 일대 포함), 200여 핵심 사업자 DB 기보유'],
        ['지자체 협업', '글로컬상권창출팀 2년 실적 (142회 현장 협업)'],
        ['B2G 수행', '누적 16.4억원 규모 대형 정부 프로젝트 완주'],
        ['창업가 발굴', '44명 발굴 실적'],
        ['행궁동행 (BID)', '사단법인 창립 총회 완료. 기획 주체'],
        ['Anchor Alliance', '정지영커피로스터즈, 진미통닭, 존앤진 등 지역 1등 기업 네트워크'],
        ['전략적 파트너십', '유통 도메인 파트너 확보 (협의 중)'],
        ['지분 구조', '박승현 60%+ / 공존공간 40%-'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §12. Ask
    # ═══════════════════════════════════════
    doc.add_heading('11. Ask — 투자 요청', level=1)

    doc.add_heading('시드 라운드: 2.5억원', level=2)

    add_table(doc, [
        ['사용처', '금액', '비중', '내용'],
        ['인건비', '1.25억원', '50%', 'Core Team 3명 급여+4대보험'],
        ['공간', '0.5억원', '20%', '보증금 3,500만 + 초기 임대료·관리비'],
        ['제품 개발 인프라', '0.35억원', '14%', 'AI API, 클라우드 서버, 개발 도구. 개발 인력은 전략적 파트너와 공동 개발'],
        ['초기 유저 확보', '0.25억원', '10%', '행궁동 100개소 온보딩, 마케팅'],
        ['예비비', '0.15억원', '6%', '비상 대응'],
    ])

    add_quote(doc, '제품 개발비가 14%로 낮은 이유: 전략적 파트너와 공동 개발 구조다.\n파트너는 슈파의 소상공인 현장 채널·데이터에 접근하고, 신도시는 개발 역량을 확보한다.\n투자금은 인프라(서버·API)에 집중하고, 개발 인력 비용은 파트너와 분담한다.\n덕분에 투자금의 절반 이상이 팀과 유저 확보에 집중된다.')

    doc.add_heading('월 고정 운영비 (Monthly Burn Rate)', level=2)

    add_table(doc, [
        ['항목', '금액 (월)'],
        ['급여 + 4대보험', '약 1,550만원'],
        ['임차료 + 관리비', '약 520만원'],
        ['AI 서비스 구독', '약 50만원'],
        ['영업 활동비', '약 100만원'],
        ['기타 관리비', '약 20만원'],
        ['합계', '약 2,240만원'],
    ], highlight_rows=[6])

    doc.add_heading('런웨이 분석 — 왜 2.5억인가', level=2)

    add_table(doc, [
        ['항목', '금액'],
        ['투자금', '2.5억원'],
        ['초기 세팅비 (보증금+장비)', '-4,545만원'],
        ['남는 운영자금', '약 2.05억원'],
    ], highlight_rows=[3])

    add_table(doc, [
        ['시나리오', '런웨이'],
        ['B2G 매출 없이 (최악)', '약 9개월'],
        ['B2G 용역 연 2억 포함', '약 15개월'],
    ])

    add_quote(doc, '2억 투자 시 런웨이 약 7개월(B2G 없이) — B2G가 한 분기만 밀려도 현금 부족 위험.\n2.5억이면 B2G 시점과 무관하게 MVP(9월) + 100개소 확보(12월)까지 안정적으로 운영 가능.')

    doc.add_heading('초기 소요 자금 (Capex)', level=2)

    add_table(doc, [
        ['항목', '금액', '비고'],
        ['임대 보증금', '3,500만원', '회수 가능 자산'],
        ['자산 매입비 (IT 장비)', '약 525만원', '노트북, 모니터 등'],
        ['예비 운영비 (첫 달)', '약 520만원', '월세/관리비'],
        ['합계', '약 4,545만원', ''],
    ], highlight_rows=[4])

    doc.add_heading('마일스톤', level=2)

    add_table(doc, [
        ['시점', '목표'],
        ['2026년 3월', '법인 설립'],
        ['2026년 9월', '슈파 MVP 출시'],
        ['2026년 12월', '행궁동 100개소 확보'],
        ['2027년 2분기', '수원 1,000개소 + 공동구매 시작 = Phase 1 완료'],
        ['2027년 1분기~', '전국 모델 확대 준비 시작'],
        ['2027년 하반기', '데이터 판매 시작'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════
    # §13. 비전
    # ═══════════════════════════════════════
    doc.add_heading('12. 비전', level=1)

    add_diagram_box(doc,
        '"기술 소외 계층 없는 디지털 경제 실현"\n\n'
        'The First City OS Provider.\n\n'
        'Phase 1              Phase 2             Phase 3\n'
        '경영 지원 +       →  데이터가         →  피지컬 AI가\n'
        '비용 절감              수익이 되고           현실에 들어온다\n\n'
        '행궁동 · 수원          전국                 라스트마일\n'
        '100 → 1,000개소        데이터 B2B           로봇 · 자동화\n\n'
        '수원 성공 모델 → 타 지자체로 City OS 복제 → 전국 단위 플랫폼\n\n'
        '         소프트웨어 → 데이터 → 하드웨어\n'
        '         디지털 세계 → 현실 세계'
    )

    # ═══════════════════════════════════════
    # §14. 기대효과
    # ═══════════════════════════════════════
    doc.add_heading('13. Expected Synergy — 기대 효과', level=1)

    add_table(doc, [
        ['관점', '효과'],
        ['투자자', '자금이 부동산에 묶이지 않고 성장·기술에만 투입 → 높은 ROE'],
        ['소상공인', '경영 비용 절감 + AI 도구 → 생존력 강화'],
        ['지역 사회', '공존 기금 → 상권 보호, 젠트리피케이션 방지'],
        ['스케일업', '수원 성공 모델 → 타 지자체 복제 → 전국 City OS'],
    ])

    # ── 마지막 페이지 ──
    doc.add_page_break()

    for _ in range(8):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('감사합니다.')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    set_run_font(run, '맑은 고딕')

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(주)신도시')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_run_font(run, '맑은 고딕')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('The First City OS Provider.')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    set_run_font(run, '맑은 고딕')

    # ── 저장 ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'IR-기획서.docx')
    doc.save(output_path)
    print(f'✅ 워드 파일 생성 완료: {output_path}')
    return output_path


# ═══════════════════════════════════════
# 헬퍼 함수들
# ═══════════════════════════════════════

def set_run_font(run, font_name):
    """한글 폰트 설정"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_quote(doc, text):
    """인용문 스타일 추가 (왼쪽 파란 선)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    # 왼쪽 보더 추가
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="4ECDC4"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.italic = True
    set_run_font(run, '맑은 고딕')


def add_diagram_box(doc, text):
    """다이어그램/코드 블록 박스"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    # 배경색 + 보더
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA" w:val="clear"/>')
    pPr.append(shd)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="D0D5DD"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="8" w:color="D0D5DD"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="D0D5DD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="8" w:color="D0D5DD"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    set_run_font(run, '맑은 고딕')


def add_table(doc, data, highlight_rows=None):
    """표 추가. data[0]은 헤더 행."""
    if highlight_rows is None:
        highlight_rows = []

    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 테이블 스타일
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            set_run_font(run, '맑은 고딕')

            if i == 0:
                # 헤더 행
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B2A4A" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shd)
            elif i in highlight_rows:
                # 강조 행
                run.bold = True
                run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F4F8" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shd)
            else:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if i % 2 == 0:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9FAFB" w:val="clear"/>')
                    cell._element.get_or_add_tcPr().append(shd)

    # 테이블 뒤 간격
    doc.add_paragraph('').paragraph_format.space_after = Pt(4)

    return table


if __name__ == '__main__':
    create_ir_document()
