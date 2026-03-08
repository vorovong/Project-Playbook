"""신도시 IR 개선 과업 지시서 생성 — 이승훈님 통합"""
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import date


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return h


def add_task(doc, number, title, details, deadline="협의 후 확정"):
    p = doc.add_paragraph()
    run = p.add_run(f"과업 {number}. {title}")
    run.bold = True
    run.font.size = Pt(12)

    for detail in details:
        doc.add_paragraph(detail, style="List Bullet")

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"마감: {deadline}")
    run2.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run2.bold = True
    doc.add_paragraph()  # spacing


def create_doc():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10.5)

    # Title
    title = doc.add_heading("신도시 IR 개선 — 과업 지시서", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Meta
    meta = doc.add_paragraph()
    meta.add_run("수신: 이승훈\n").bold = True
    meta.add_run("발신: 박승현 대표\n")
    meta.add_run(f"날짜: {date.today().strftime('%Y년 %m월 %d일')}\n")
    meta.add_run("관련: 신도시 IR 기획서 v0.2 → FI 피드백 대응 개선")
    meta.paragraph_format.space_after = Pt(12)

    # Background
    add_heading_styled(doc, "배경", level=1)
    doc.add_paragraph(
        '투자 전문가(스마트스터디벤처스 등)로부터 "비즈니스 모델 및 수익 시나리오 빈약"이라는 '
        "피드백을 받았습니다. 프론티어 분석(Toast, GrubMarket, 마켓보로, Choco, Slice) 결과, "
        "핵심 개선 영역이 도출되었습니다."
    )

    items = [
        "Unit Economics(단위 경제학) 부재 → 가게 1개 단위 수익 구조 필요",
        '"무료 → 유료" 전환 트리거 불명확',
        "Lock-in(이탈 방지) 메커니즘 부족",
        "TAM이 Top-down → Bottom-up 재구성 필요",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # === Section A: 현장 데이터 수집 ===
    add_heading_styled(doc, "A. 현장 데이터 수집 (최우선)", level=1)

    add_task(
        doc,
        1,
        "행궁동 현장 데이터 수집 (Unit Economics 실측용)",
        [
            "IR 보강의 가장 중요한 기초 자료입니다",
            "행궁동 내 5~10개소 대상으로 아래 항목을 인터뷰/조사해주세요:",
            "  - 월 식자재/소모품 구매액 (총액)",
            "  - 주요 구매 품목 Top 5 (품목명 + 월 구매액)",
            "  - 현재 유통 경로 (포터 아저씨, 직접 구매, 마트 등) + 각 경로의 단가",
            "  - 현재 경영 관리 방식 (수기, 엑셀, 앱 등)",
            "  - 슈파 같은 서비스가 있다면 쓸 의향이 있는지",
            "업종 분포: 식음(카페, 식당, 주점) 위주 + 비식음 1~2개소 포함",
            "형식: 가게별로 정리. 엑셀이면 가장 좋음",
        ],
    )

    add_task(
        doc,
        2,
        "구매의향서 확보 (10개소 이상)",
        [
            "슈파 공동구매 서비스 이용 의향이 있는 사업장의 구매의향서를 확보해주세요",
            "목표: 최소 10개소",
            "의향서 내용: 사업장명, 대표자명, 업종, 예상 월 구매액, 서명/날인",
            "형식은 간단해도 됩니다 — 핵심은 '실제 수요가 있다'는 증거",
            "이 자료가 IR에서 투자자를 설득하는 핵심 무기가 됩니다",
            "구매의향서 양식이 필요하면 말씀해주세요. 만들어드리겠습니다",
        ],
    )

    add_task(
        doc,
        3,
        "유통 파트너 단가 조사",
        [
            "CJ 프레시웨이 등 주요 식자재 유통사에 공급 단가를 문의해주세요",
            "비교 대상: 현재 행궁동 소상공인들이 쓰는 로컬 유통(포터 아저씨 등)",
            "확인할 것: 최소 물량 조건, 단가, 배송 조건, 결제 조건",
            "목표: '슈파 공동구매로 묶으면 이만큼 싸진다'를 숫자로 증명",
            "유통사 2~3곳 비교하면 충분합니다",
        ],
    )

    add_task(
        doc,
        4,
        "Bottom-up TAM 검증 데이터",
        [
            "행궁동 상권 내 업종별 사업체 수를 실측 확인해주세요",
            "기존 데이터: 장안동·신풍동 기준 449개 (식음 54%)",
            "구천동·남문 일대 포함 약 2,400개 사업자 (현장 파악)",
            "검증할 것: 이 숫자가 정확한지, 업종별 분포가 어떤지",
            "가능하면 업종별(카페/식당/주점/소매/서비스 등) 개수 정리",
            "과업 1과 동시 진행 가능합니다",
        ],
    )

    # === Section B: 전략/설계 ===
    add_heading_styled(doc, "B. 전략/설계 (현장 데이터와 병행)", level=1)

    add_task(
        doc,
        5,
        "Unit Economics 구조 설계",
        [
            "가게 1개 단위의 경제 구조를 산출할 수 있는 공식을 만들어주세요",
            "필요 지표: CAC(고객 획득 비용), LTV(고객 생애 가치), ARPU(가게당 월 매출), CAC Payback(회수 기간)",
            "과업 1에서 수집한 현장 데이터를 넣으면 바로 숫자가 나오는 구조",
            "참고: Toast의 시드 라운드 핵심 숫자 — CAC Payback 14개월, 순유지율 150%, 연간 이탈 3%",
            "이 3가지에 대응하는 슈파 버전 지표를 설계해주세요",
        ],
    )

    add_task(
        doc,
        6,
        "Lock-in 메커니즘 설계",
        [
            '현재 IR에서 "왜 고객이 못 떠나는가?"에 대한 구조적 답이 없습니다',
            "슈파 서비스의 전환 비용(Switching Cost)을 만드는 구조를 설계해주세요",
            "참고 사례: 마켓보로 — 거래 데이터 축적으로 떠나면 거래처 정보 소실",
            "참고 사례: Toast — 하드웨어 설치로 물리적 전환 비용",
            "제안 방향: AI 자동 발주 학습 데이터가 가게에 종속 → 다른 서비스로 이전 불가",
            "데이터 축적이 곧 서비스 품질 향상 → 오래 쓸수록 정확도 상승 → 떠날 이유 감소",
        ],
    )

    add_task(
        doc,
        7,
        "Phase 전환 트리거 정의",
        [
            "현재 IR의 Phase 1→2→3 전환이 추상적입니다. 구체적 전환 조건을 정의해주세요",
            "Phase 1→2 전환 조건 예시: 가맹점 N개 + 월 거래액 M원 + AI 발주 정확도 90%+",
            "각 조건의 구체적 숫자를 제안해주세요 (근거 포함)",
            "Phase 2 깊이 확장 기능 검토: 결제/정산 대행, 세무 대행, 소상공인 대출 중개",
            "참고: Toast는 POS 설치 수 → 결제 볼륨 → 소프트웨어 부착률 순으로 전환 트리거를 설정",
            "⚠️ Phase 3(피지컬 AI)는 FI용 IR에서 축소 예정 — 장기 비전 1장으로",
        ],
    )

    # === 우선순위 ===
    add_heading_styled(doc, "우선순위 정리", level=1)
    doc.add_paragraph("① 과업 1 (현장 데이터) — 모든 숫자의 기초. 최우선")
    doc.add_paragraph("② 과업 2 (구매의향서) — IR 핵심 무기. 과업 1과 동시 진행")
    doc.add_paragraph("③ 과업 3 (유통 단가) — 가격 경쟁력 증명")
    doc.add_paragraph("④ 과업 4 (TAM 검증) — 과업 1과 동시 진행 가능")
    doc.add_paragraph("⑤ 과업 5~7 (전략/설계) — 현장 데이터 들어오는 대로 병행")

    # === 참고 사항 ===
    add_heading_styled(doc, "참고 사항", level=1)
    doc.add_paragraph(
        "현장 데이터(A 영역)가 IR 숫자의 기초가 됩니다. "
        "완벽하지 않아도 괜찮습니다. 5개소 데이터만 있어도 시작할 수 있습니다."
    )
    doc.add_paragraph(
        "결과물은 문서, 엑셀, 슬라이드 어떤 형태든 괜찮습니다. "
        "핵심은 '현장의 실제 숫자'와 '논리 구조'입니다."
    )
    doc.add_paragraph("궁금한 점은 바로 연락 주세요.")

    return doc


if __name__ == "__main__":
    import os
    out_dir = r"C:\Users\Administrator\projects\Project_Playbook\mini-projects\박승현\신도시-IR"

    doc = create_doc()
    path = os.path.join(out_dir, "과업지시-이승훈.docx")
    doc.save(path)
    print(f"생성: {path}")
