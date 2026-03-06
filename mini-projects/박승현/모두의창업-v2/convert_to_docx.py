"""
수원형-모두의창업-v2-내부검토.md → .docx 변환 스크립트
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════
# 헬퍼 함수들
# ═══════════════════════════════════════

def set_run_font(run, font_name='맑은 고딕'):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_run(p, text, bold=False, italic=False, size=Pt(10.5), color=RGBColor(0x33, 0x33, 0x33)):
    run = p.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = '맑은 고딕'
    set_run_font(run)
    return run


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        set_run_font(run)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10.5), color=RGBColor(0x33, 0x33, 0x33)):
    p = doc.add_paragraph()
    add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="4ECDC4"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.italic = True
    set_run_font(run)
    return p


def add_diagram_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA" w:val="clear"/>')
    pPr.append(shd)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="D0D5DD"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="8" w:color="D0D5DD"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="D0D5DD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="8" w:color="D0D5DD"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    set_run_font(run)
    return p


def add_table(doc, data, highlight_rows=None):
    if highlight_rows is None:
        highlight_rows = []
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            set_run_font(run)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B2A4A" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shd)
            elif i in highlight_rows:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F4F8" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shd)
            else:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if i % 2 == 0:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9FAFB" w:val="clear"/>')
                    cell._element.get_or_add_tcPr().append(shd)
    doc.add_paragraph('').paragraph_format.space_after = Pt(4)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        set_run_font(run)
    return p


# ═══════════════════════════════════════
# 메인 문서 생성
# ═══════════════════════════════════════

def create_document():
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 기본 스타일
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 제목 스타일
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = '맑은 고딕'
        h_font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        h_rPr = h_style.element.get_or_add_rPr()
        h_rFonts = h_rPr.find(qn('w:rFonts'))
        if h_rFonts is None:
            h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
            h_rPr.append(h_rFonts)
        else:
            h_rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(24)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(12)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(18)
    doc.styles['Heading 2'].paragraph_format.space_after = Pt(8)
    doc.styles['Heading 3'].font.size = Pt(12)
    doc.styles['Heading 3'].paragraph_format.space_before = Pt(12)
    doc.styles['Heading 3'].paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════
    # 표지
    # ═══════════════════════════════════
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "수원형 '모두의창업' 생태계 구축", bold=True, size=Pt(28), color=RGBColor(0x1B, 0x2A, 0x4A))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '사업계획서 (내부 검토용)', size=Pt(14), color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "이재명의 판교, 정원오의 성수동, 이재준의 '행궁동'", bold=True, size=Pt(13), color=RGBColor(0x4E, 0xCD, 0xC4))

    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026. 03. | (주)공존공간', size=Pt(12), color=RGBColor(0x66, 0x66, 0x66))

    doc.add_page_break()

    # ═══════════════════════════════════
    # 1. 한 장 요약
    # ═══════════════════════════════════
    add_heading(doc, '1. 한 장 요약', level=1)

    add_table(doc, [
        ['항목', '내용'],
        ['사업명', "수원형 '모두의창업' 생태계 구축"],
        ['총 예산', '10억 원 (수원시 + 중기부 연계)'],
        ['기간', '2026. 07 ~ 2027. 06 (1년차)'],
        ['핵심 사업', '4대 축 — 지역밀착AC / 중기부연계 / 시자원연계 / 락성연(축제)'],
        ['대상지', '행궁동 글로컬상권 + 남문시장 백년시장 → 수원 전역 확대'],
        ['운영', '민간 CMO 사단법인 (관광+상권+창업 통합)'],
    ])

    add_quote(doc, '2024~2025년 글로컬상권에서 이미 증명한 모델을 공식 정책으로 편입하고,\n수원화성문화제와 결합하여 전국 최초 테크+로컬 통합 창업 생태계를 구축한다.')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 2. 왜 지금인가
    # ═══════════════════════════════════
    add_heading(doc, '2. 왜 지금인가 — 세 개의 파도', level=1)

    add_diagram_box(doc,
        '2026년, 세 파도가 동시에 온다\n\n'
        '[파도 1] 국가창업시대         [파도 2] 수원방문의해       [파도 3] 소상공인 AI 전환\n'
        '  모두의창업 4단계              관광객 1,500만 목표         AI활용지원 114~194억\n'
        '  창업예산 3.4조               수원화성문화제 112만+       신규 예산\n'
        '  벤처펀드 4.4조               3대 축제 체계 구축\n\n'
        '         ──── 세 파도를 한 번에 탈 도시 = 수원 ────'
    )

    add_heading(doc, '수원이 가진 것', level=2)

    add_table(doc, [
        ['테크 자산', '로컬 자산', '관광 자산'],
        ['환상형 혁신클러스터', '글로컬상권 (전국 3곳 중 1곳)', '수원방문의해 2026-27'],
        ['새빛펀드 7,600억', '백년시장 (40억, 8개소)', '세계 3대 축제도시 추진'],
        ['AC 2개 운용사 신규', '행궁동행 BID', '112만 방문객 (2025)'],
    ])

    add_para(doc, '전국에서 테크와 로컬을 균형 있게 추진하며, 축제라는 무대까지 가진 도시 = 수원뿐이다.', bold=True)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 3. 이미 증명된 것
    # ═══════════════════════════════════
    add_heading(doc, '3. 이미 증명된 것 — 글로컬상권 2년의 실적', level=1)

    add_para(doc, '이 사업은 구상이 아니라 실적 위에 세우는 것이다.', bold=True)

    add_heading(doc, '3-1. 글로컬상권창출사업 핵심 성과 (2024.06~2025.12)', level=2)

    add_table(doc, [
        ['지표', 'Before', 'After', '변화'],
        ['외국인 소비액', '14억', '29억', '+107%'],
        ['소상공인 협업', '20회', '142회', '+710%'],
        ['창업가 발굴', '0명', '44명', '신규'],
        ['1인당 소비액', '9,333원', '13,810원', '+48%'],
        ['대기업 협업', '없음', '롯데, 오뚜기, 경희대, 아주대', '비예산 자생'],
    ])

    add_para(doc, '방문객 40% 증가에 소비액 107% 증가 = 양이 아니라 질이 바뀌었다.', bold=True)

    add_heading(doc, '3-2. 깍페스티벌 — 축제 모델의 파일럿', level=2)

    add_table(doc, [
        ['', '깍페스티벌(상) 2025 상반기', '깍페스티벌(하) 2025.09'],
        ['참여자', '355명 (사전신청 97명)', '목표 100개소 소상공인 참여'],
        ['온라인 유입', '13,292명', '다국어 콘텐츠 50건+'],
        ['협업', '65개소 쿠폰 발행', '60개소+ 팝업 모집'],
        ['콘셉트', 'IF 정조 NOW', '수원 깍쟁이의 기업가정신'],
    ])

    add_heading(doc, '3-3. 락성연 — 이미 실행된 축제 프로토타입', level=2)

    add_para(doc, "2025년 깍페스티벌(하)에서 '락(樂)성연'이라는 이름으로 축제를 실행했다.")

    add_table(doc, [
        ['항목', '내용'],
        ['개막식', '9/26(금) 17:30~22:00, 롯데제과수원영업소'],
        ['프로그램', '랜덤플레이댄스 → 개막식(시장·국회의원 참석) → 시민장기자랑 → 수월래'],
        ['공연', '9/27(토) 16:00~22:30, DJ 5팀 + 밴드 7팀'],
        ['포지셔닝', "수원화성문화제 '비공식 사전행사' 실험"],
        ['주최', '중기부, 소진공, 수원시, 수원도시재단, 공존공간'],
        ['연계', '수원 AI써밋, 경희대 국제세미나, 수원화성문화제 동기간'],
    ])

    add_quote(doc, '락성연은 "새로 만들 것"이 아니다.\n2025년에 비공식으로 실험한 것을, 2026년에 공식 편입하고 스케일업하는 것이다.')

    add_heading(doc, '3-4. 현장에서 발견한 하나의 인사이트', level=2)

    add_quote(doc, '"교육하고 컨설팅하는 간접 지원은 한계가 있다.\n소상공인에게는 파트너로서 직접 효과를 주는 구조가 필요하다."')

    add_para(doc, "소상공인을 '보호 대상'이 아니라 '도전하고 성장하는 주체'로 본다. 이것이 수원형 모두의창업의 설계 원리다.")

    doc.add_page_break()

    # ═══════════════════════════════════
    # 4. 사업 구조
    # ═══════════════════════════════════
    add_heading(doc, '4. 사업 구조 — 4대 축 + 통합 운영', level=1)

    add_diagram_box(doc,
        '                    ★ 락성연 (4축)\n'
        '                    수원크리에이티브페스티벌\n'
        '                    2025 비공식 실험 → 2026 공식 편입\n'
        '                           ▲\n'
        '              ┌────────────┼────────────┐\n'
        '              │            │            │\n'
        '         1축 지역밀착AC   2축 중기부연계   3축 시자원연계\n'
        '         발굴·보육·투자    정책자금·AI     인프라·네트워크\n'
        '              │            │            │\n'
        '              └────────────┼────────────┘\n'
        '                           ▼\n'
        '                    민간 CMO 사단법인\n'
        '                    (관광+상권+창업 통합)'
    )

    add_para(doc, '1~3축이 창업가를 발굴하고, 투자를 연결하고, 자원을 모은다.\n4축 락성연이 그 모든 것을 세계에 보여주는 무대가 된다.', bold=True)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 5-1. 1축
    # ═══════════════════════════════════
    add_heading(doc, '5. 4대 축 상세 사업계획', level=1)

    add_heading(doc, '5-1. 1축: 지역밀착 AC — 로컬벤처스튜디오', level=2)

    add_para(doc, '상권기획자가 현장에서 직접 발굴 → 보육 → 투자연계를 일괄 수행.', bold=True)
    add_para(doc, '기존 "공모 → 선정 → 지원" 방식이 아니라, 2년간 200+ 사업자와 직접 협업한 DB와 관계망 기반.')

    add_table(doc, [
        ['단계', '내용', '담당', '일정'],
        ['발굴', '행궁동·남문시장 현장 스카우팅, 글로컬상권 2년 DB 활용', '민간(CMO법인) + 시', '2026.07~ 상시'],
        ['보육', '실전창업학교 운영, 대학 협업(경희대·아주대 RISE)', '민간 + 대학', '2026.08~'],
        ['투자연계', '중기부 AC 등록 → LIPS I·II 연계, 새빛펀드 AC 트랙', '민간 + 시 + 중기부', '2026.10~'],
        ['무대', '락성연 피칭 기회, 트레이드쇼 출전', '민간(CMO법인)', '2026.10'],
    ])

    add_heading(doc, 'LIPS 연계 (로컬 버전 TIPS)', level=3)

    add_table(doc, [
        ['구분', 'LIPS I', 'LIPS II', '합계'],
        ['성격', '매칭융자', '투자연계 보조금', ''],
        ['지원 한도', '최대 5억', '최대 2억', '최대 9억/건'],
    ])

    add_para(doc, '1년차 KPI:', bold=True)
    add_bullet(doc, '발굴 50명+')
    add_bullet(doc, '실전창업학교 2기 운영 (수료자 40명+)')
    add_bullet(doc, 'LIPS 연계 신청 5건+')
    add_bullet(doc, '중기부 AC 등록 완료')

    add_para(doc, '예산: 2.5억 원', bold=True)
    add_table(doc, [
        ['항목', '금액', '비고'],
        ['현장 스카우팅·DB 운영', '0.5억', '인건비 포함'],
        ['실전창업학교 2기 운영', '1.0억', '강사·공간·프로그램'],
        ['AC 등록·LIPS 연계 행정', '0.3억', '법무·회계'],
        ['대학 협업 프로그램', '0.5억', '경희대·아주대 RISE'],
        ['예비비', '0.2억', ''],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 5-2. 2축
    # ═══════════════════════════════════
    add_heading(doc, '5-2. 2축: 중기부 연계 — 국가 정책 직접 연동', level=2)

    add_para(doc, '중기부 담당자가 이미 긍정 검토 중. 수원화성문화제 112만 관객 앞에서 모두의 창업 성과를 보여줄 무대는 전국 어디에도 없다.', bold=True)

    add_table(doc, [
        ['중기부 사업', '수원 연계 방안', '예상 효과'],
        ['모두의 창업 4단계', '락성연을 경기 권역 본선 유치 무대로', '수원 = 모두의창업 대표 도시'],
        ['혁신 소상공인 AI활용 (114~144억)', '수원을 AI 전환 시범 도시 추진', '전국 1호 모델'],
        ['AI도우미 지원 (50억)', '수원 소상공인 대상 시범 운영', '현장 데이터 확보'],
        ['로컬크리에이터 육성', '1축 발굴 크리에이터 연계', '중복 없는 파이프라인'],
        ['글로컬상권 사업', '행궁동 → 남문시장 → 수원 전역 확대', '글로컬 모델의 전국 표준화'],
    ])

    add_para(doc, '예산: 1.0억 원 (시 매칭분, 중기부 예산은 별도)', bold=True)
    add_table(doc, [
        ['항목', '금액', '비고'],
        ['중기부 협의·사업 기획', '0.3억', ''],
        ['AI바우처 시 매칭·교육', '0.5억', '소상공인 AI 교육 포함'],
        ['예비비', '0.2억', ''],
    ])

    # ═══════════════════════════════════
    # 5-3. 3축
    # ═══════════════════════════════════
    add_heading(doc, '5-3. 3축: 시 자원 연계 — 수원만의 자산 활용', level=2)

    add_table(doc, [
        ['자원', '연계 방안', '담당'],
        ['새빛펀드 7,600억', 'AC 로컬창업 트랙 배정 → 1축 직접 연결', '지역경제과'],
        ['수원페이 411억 인센티브', '락성연 기간 트레이드쇼 결제수단 연동', '경제정책과'],
        ['경희대·아주대 (RISE)', '로컬콘텐츠중점대학, 보육 협업', '민간+대학'],
        ['삼성·신세계·롯데·오뚜기', '글로컬상권에서 이미 협업 구축. 트레이드쇼 바이어·스폰서', '민간'],
        ['KB금융-수원도시재단 데이터', '소상공인 데이터 기반 정책 수립 (이미 추진 중)', '도시재단'],
        ['수원방문의해 마케팅', '락성연을 방문의해 핵심 이벤트로 편입', '관광과'],
    ])

    add_para(doc, '예산: 1.0억 원', bold=True)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 5-4. 4축 락성연
    # ═══════════════════════════════════
    add_heading(doc, '5-4. 4축: 락성연 — 수원크리에이티브페스티벌', level=2)

    add_para(doc, '수원화성문화제 기간 중 공식 병행 개최하는 테크+로컬 융합 축제.', bold=True)
    add_para(doc, '2025년 비공식 사전행사로 실험 완료. 이제 공식 편입하여 스케일업한다.')

    add_heading(doc, '2025 → 2026 진화 비교', level=3)

    add_table(doc, [
        ['', '2025 (비공식 실험)', '2026 (공식 편입)'],
        ['위상', '깍페스티벌 내 프로그램', '수원화성문화제 공식 연계 행사'],
        ['규모', '개막식+공연 2일', '3~4일, 5개 트랙'],
        ['주최', '공존공간 (글로컬상권사업 내)', '수원시·중기부 공동'],
        ['대기업 협업', '롯데제과 장소 협업', '대기업·테크기업이 오디션 심사·스폰서·바이어로 참여'],
        ['참여', '소상공인·주민 중심', '+ 스타트업·투자자·바이어'],
        ['예산', '글로컬상권 사업비 내', '별도 10억 사업 예산'],
    ])

    add_heading(doc, '프로그램 구성 (3~4일)', level=3)

    add_table(doc, [
        ['트랙', '내용', '일정(안)'],
        ['락성연 개막제', '랜덤플레이댄스 + 개막식 + 시민장기자랑 + 수월래\n(2025 검증 프로그램 확대)', 'Day 1 저녁'],
        ['모두의창업 오디션', '로컬크리에이터·소상공인·스타트업이\n대기업·테크기업·투자자 앞에서 발표.\n심사위원: 롯데·오뚜기·삼성·신세계 + AI·테크기업\n→ 현장 협업·투자·유통 연결', 'Day 2 오전~오후'],
        ['트레이드쇼', '로컬 브랜드 전시·판매 + 대기업 바이어 초청.\n수원페이 결제 연동. 2025 롯데제과 협업 확대', 'Day 2~3 전일'],
        ['AI 데모데이', '소상공인 맞춤 AI 솔루션 시연.\n테크기업이 소상공인에게 직접 솔루션 제안', 'Day 2 오후'],
        ['락성연 공연', 'DJ + 밴드 라이브\n(2025 2스테이지 → 야외 메인+서브 3스테이지)', 'Day 2~3 저녁'],
        ['네트워킹 나이트', '수원화성 미디어아트 연계 야간 네트워킹', 'Day 3 저녁'],
        ['컨퍼런스', '로컬경제·AI·소상공인 미래 강연. 국내외 연사', 'Day 3 오전'],
    ])

    add_heading(doc, 'SXSW와의 차별점', level=3)

    add_table(doc, [
        ['', 'SXSW (오스틴)', '락성연 (수원)'],
        ['시작', '1987 음악 축제', '2025 깍페스티벌 내 실험'],
        ['중심', '테크 스타트업', '테크 + 로컬 소상공인'],
        ['기반', '도시 자체', '유네스코 세계유산'],
        ['관광 연계', '부수적', '수원방문의해 직접 연동'],
        ['독자성', '테크 중심 축제는 세계에 많음', '테크+로컬 융합은 세계에 없음'],
    ])

    add_heading(doc, '중기부 "모두의 창업" 4단계 직접 연결', level=3)

    add_table(doc, [
        ['모두의 창업 단계', '락성연 연계'],
        ['1단계: 창업가 5,000명 발굴 (1인 200만원)', '수원 권역에서 로컬 창업가 발굴 (1축)'],
        ['2단계: 100명 루키 선발 (1인 1억원)', '락성연 피칭이 경기 권역 본선 무대'],
        ['3단계: 전국 경진대회 (우승 10억원)', '락성연을 전국 경진대회 개최지로 유치'],
        ['4단계: 5,000억 창업 펀드', '새빛펀드 AC 연계, 투자 대상 발굴 현장'],
    ])

    add_para(doc, '1년차 KPI:', bold=True)
    add_bullet(doc, '수원화성문화제 공식 연계 행사 지정')
    add_bullet(doc, '3~4일 운영, 참여 소상공인 150개소+')
    add_bullet(doc, '모두의창업 오디션 참가 30팀+, 대기업·테크기업 심사위원 10명+')
    add_bullet(doc, '오디션 통해 협업·투자·유통 연결 10건+')
    add_bullet(doc, '참여 관객 2만명+')

    add_para(doc, '예산: 4.5억 원', bold=True)
    add_table(doc, [
        ['항목', '금액', '비고'],
        ['락성연 개막제 + 공연', '1.5억', '무대·음향·출연료·안전관리'],
        ['트레이드쇼 설치·운영', '1.0억', '부스·물류·수원페이 연동'],
        ['모두의창업 오디션 + AI 데모데이', '0.5억', '대기업·테크기업 심사위원단·상금·운영'],
        ['컨퍼런스 + 네트워킹', '0.5억', '연사·장소·케이터링'],
        ['홍보·마케팅', '0.7억', '옥외광고, SNS, 다국어 콘텐츠'],
        ['예비비', '0.3억', ''],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 5-5. CMO
    # ═══════════════════════════════════
    add_heading(doc, '5-5. 통합 운영: 민간 CMO 사단법인', level=2)

    add_table(doc, [
        ['항목', '내용'],
        ['역할', '4대 축 통합 조율 + 관광·상권·창업 허브 + 사업 기획의 지속성 확보'],
        ['모델', 'NYC & Company (뉴욕시티앤컴퍼니) — 시 예산 + 민간 운영'],
        ['왜 민간인가', '행정이 바뀌어도 남는 구조. 현장 속도. 장기 축적.'],
        ['시작점', '행궁동행 사단법인 (이미 설계 중)'],
    ])

    add_heading(doc, 'CMO 자생 수익모델 — 2년차부터 기획비용 자체 조달', level=3)

    add_para(doc, '1년차는 수원시·중기부 재원으로 운영하되, 2년차부터 유료 프로그램을 통해 CMO가 자체 기획비용을 확보한다.')

    add_table(doc, [
        ['수익원', '내용', '시작 시점'],
        ['유료 컨퍼런스', '락성연 컨퍼런스 유료 티켓 (일반 3만원 / VIP 10만원)', 'Y2'],
        ['트레이드쇼 부스비', '출전 브랜드 부스 참가비 (50~100만원/부스)', 'Y2'],
        ['체류형 유료 프로그램', '로컬 러닝·투어·워크숍 참가비', 'Y2'],
        ['스폰서십 패키지', '대기업·브랜드 스폰서 패키지', 'Y2'],
        ['피칭 참가비', '스타트업·로컬브랜드 피칭 참가 수수료', 'Y3'],
        ['콘텐츠 라이선스', '상권 데이터·콘텐츠 타 지역 판매', 'Y3'],
    ])

    add_quote(doc, 'CMO가 기획비용을 자체 조달할 수 있게 되면,\n행정 교체와 무관하게 사업이 지속된다.')

    add_para(doc, '1년차 예산: 1.0억 원', bold=True)
    add_table(doc, [
        ['항목', '금액', '비고'],
        ['사단법인 설립·운영', '0.5억', '법인 설립, 초기 인건비'],
        ['4대 축 조율·기획', '0.3억', ''],
        ['2년차 유료 프로그램 기획', '0.1억', '수익모델 설계·시장조사'],
        ['예비비', '0.1억', ''],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 6. 예산 총괄
    # ═══════════════════════════════════
    add_heading(doc, '6. 예산 총괄 (10억 원)', level=1)

    add_table(doc, [
        ['축', '사업', '예산', '비중'],
        ['1축', '지역밀착 AC (발굴·보육·투자)', '2.5억', '25%'],
        ['2축', '중기부 연계 (AI·정책자금)', '1.0억', '10%'],
        ['3축', '시 자원 연계 (새빛펀드·대학·기업)', '1.0억', '10%'],
        ['4축', '락성연 (축제 플랫폼)', '4.5억', '45%'],
        ['통합', '민간 CMO 사단법인', '1.0억', '10%'],
        ['합계', '', '10.0억', '100%'],
    ], highlight_rows=[5, 6])

    add_heading(doc, '재원 구성(안)', level=2)

    add_para(doc, '이 사업은 첫해이므로 AC가 자체적으로 예산을 만드는 것은 무리가 있다.\n수원시와 중기부가 재원을 책정하고, AC·CMO법인은 실행 주체로서 사업을 운영한다.', bold=True)

    add_table(doc, [
        ['재원', '금액', '비고'],
        ['수원시 예산', '5~6억', '기존 축제 예산 일부 + 지역경제 예산'],
        ['중기부 연계', '3~4억', '모두의창업·글로컬상권·AI바우처 등'],
        ['민간 협찬', '1억', '대기업 스폰서(롯데·오뚜기 등 기존 협업사)'],
    ])

    add_quote(doc, '역할 구분: 수원시·중기부 = 재원 출처 + 정책 방향 / AC·CMO법인 = 기획·운영·실행')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 7. 실행 일정
    # ═══════════════════════════════════
    add_heading(doc, '7. 실행 일정', level=1)

    add_table(doc, [
        ['시기', '내용', '담당'],
        ['2026.03~04', '내부 검토 완료 → 시장 재보고', '공존공간·지역경제과'],
        ['2026.04~05', '중기부 담당자 미팅, 모두의창업 경기권역 유치 협의', '공존공간·중기부'],
        ['2026.05~06', '사단법인 설립, 사업 예산 확정, 관광과 락성연 공식 편입 협의', '지역경제과·관광과'],
        ['2026.07', '1축 현장 스카우팅 시작, 실전창업학교 모집', 'CMO법인'],
        ['2026.08', '2축 AI바우처 수원 매칭 시작, 3축 대학·기업 협업 착수', 'CMO법인·시'],
        ['2026.09', '락성연 사전 프로그램 (깍페스티벌 연계)', 'CMO법인'],
        ['2026.10', '★ 락성연 본행사 (수원화성문화제 기간)', '전체'],
        ['2026.11~12', '성과 분석, 2년차 계획 수립', 'CMO법인·시'],
        ['2027', '2년차: 규모 확대 + 수원방문의해 연계 강화 + 유료 프로그램 시작', ''],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 8. 역할 분장
    # ═══════════════════════════════════
    add_heading(doc, '8. 관(官)-민(民)-중기부 역할 분장', level=1)

    add_table(doc, [
        ['', '발굴', '보육', '투자연계', '무대(락성연)'],
        ['관(官)\n수원시·도시재단', '정책 방향·예산', '인프라·공간', '새빛펀드 AC 트랙', '화성문화제 내\n공간·예산'],
        ['민(民)\nCMO 사단법인', '현장 스카우팅', '대학·기업 협업', 'LIPS I·II 연계', '락성연 기획·운영'],
        ['중기부', '모두의창업\n발굴 연계', 'AI바우처·\n디지털전환', '펀드·정책자금', '경진대회·\n홍보 무대'],
    ])

    add_para(doc, '대상지: 행궁동 글로컬상권 + 남문시장 백년시장(8개소) → 수원 전역', bold=True)

    # ═══════════════════════════════════
    # 9. 성장 로드맵
    # ═══════════════════════════════════
    add_heading(doc, '9. 성장 로드맵', level=1)

    add_table(doc, [
        ['연차', '목표', '핵심 마일스톤'],
        ['Y1 (2026)', '공식 편입 + 파일럿 스케일업', '락성연 3~4일, 화성문화제 공식 연계, CMO법인 설립'],
        ['Y2 (2027)', '본격 성장', '수원방문의해 연계 극대화, 중기부 공동 주관, 유료 프로그램 시작'],
        ['Y3 (2028)', '독립 브랜드화', '락성연 단독 BI, 5일+, Slush\'D 모델 검토'],
        ['Y5 (2030)', '아시아 대표', '아시아 로컬+테크 대표 축제, 경제효과 1,000억+'],
    ])

    add_heading(doc, 'SXSW 40년의 궤적 참고', level=3)

    add_table(doc, [
        ['', 'SXSW', '락성연 (목표)'],
        ['경제효과', '$3.77억 (약 5,000억원)', '1,000억원+ (5년차)'],
        ['도시 브랜드', '"오스틴 = 창업 도시"', '"수원 = 테크+로컬 창업 도시"'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 10. 기대 효과
    # ═══════════════════════════════════
    add_heading(doc, '10. 기대 효과', level=1)

    add_heading(doc, '정책 선점', level=2)

    add_table(doc, [
        ['선점 항목', '의미'],
        ["'모두의 창업' 로컬 분야 전국 1호 모델", '중기부 정책 성과의 대표 사례'],
        ['로컬벤처스튜디오 최초 실증', 'AC 등록 + 로컬 특화의 전국 첫 사례'],
        ['테크+로컬 통합 축제 세계 최초', 'SXSW에도 옥토버페스트에도 없는 융합'],
        ['소상공인 AI 전환 시범 도시 1호', '전국 확산 모델의 원형'],
    ])

    add_heading(doc, '실질 성과 (1년차)', level=2)

    add_table(doc, [
        ['성과', '목표'],
        ['로컬 창업가 발굴', '50명+'],
        ['LIPS·새빛펀드 투자 연결', '5건+'],
        ['락성연 참여 관객', '2만명+'],
        ['트레이드쇼 참여 소상공인', '150개소+'],
        ['상권 확산', '행궁동 → 남문시장'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 11. 협조 요청
    # ═══════════════════════════════════
    add_heading(doc, '11. 협조 요청 사항', level=1)

    add_table(doc, [
        ['#', '협조 요청 내용', '담당', '연계'],
        ['1', "행궁동 글로컬상권을 '수원형 모두의창업' 핵심 대상지로 공식 지정", '지역경제과', '1축'],
        ['2', '새빛펀드 2차 AC 트랙 내 로컬창업 전용 배정 검토', '지역경제과', '1축+3축'],
        ['3', "수원화성문화제 내 '락성연' 공식 연계 행사 편입 및 중기부 공동 주관 검토", '관광과', '4축'],
        ['4', "수원을 중기부 '소상공인 AI 전환 시범 도시'로 추진", '지역경제과', '2축'],
        ['5', '민간 CMO 사단법인 설립 지원 검토', '지역경제과', '통합'],
        ['6', "중기부 '모두의 창업' 경기 권역 본선을 락성연에서 개최 협의", '지역경제과', '2축+4축'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════
    # 부록
    # ═══════════════════════════════════
    add_heading(doc, '부록: 글로컬상권 전국 현황', level=1)

    add_table(doc, [
        ['도시', '운영사', '특징'],
        ['수원', '공존공간', '행궁동 — 세계유산+관광+소상공인'],
        ['전주', '크립톤', '글로컬 소셜클럽'],
        ['통영', '로컬스티치', '바다의 땅 통영'],
    ])

    add_quote(doc, '중기부 장관: "세계인이 찾는 글로컬 상권, 지역마다 한 개 이상 만들겠다"')

    add_para(doc, '수원이 이 3곳 중 가장 먼저 "글로컬상권 + 창업 생태계 + 축제 플랫폼"을 통합한 모델을 만들면, 전국 확산의 표준이 된다.')

    # ═══════════════════════════════════
    # 마지막 페이지
    # ═══════════════════════════════════
    doc.add_page_break()

    for _ in range(8):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '이미 증명된 것을 공식화하고, 스케일업한다.', bold=True, size=Pt(14), color=RGBColor(0x1B, 0x2A, 0x4A))

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '테크와 로컬의 균형 — 전국 최초 통합 창업 생태계.', bold=True, size=Pt(16), color=RGBColor(0x4E, 0xCD, 0xC4))

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '수원이 먼저 하면, 전국이 따라온다.', bold=True, size=Pt(20), color=RGBColor(0x1B, 0x2A, 0x4A))

    for _ in range(4):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '(주)공존공간', size=Pt(12), color=RGBColor(0x66, 0x66, 0x66))

    # 저장
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '수원형-모두의창업-v2-내부검토.docx')
    doc.save(output_path)
    print(f'워드 파일 생성 완료: {output_path}')
    return output_path


if __name__ == '__main__':
    create_document()
