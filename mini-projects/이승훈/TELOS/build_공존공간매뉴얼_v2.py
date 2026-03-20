"""공존공간 통합 운영 매뉴얼 — Word 문서 생성 스크립트"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 경로 설정 ──
MD_PATH = Path(__file__).parent / "공존공간-운영매뉴얼.md"
OUT_PATH = Path.home() / "Desktop" / "공존공간_통합운영매뉴얼.docx"

# ── 스타일 상수 ──
FONT_BODY = "맑은 고딕"
COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)   # 남색
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # 파란색
COLOR_GRAY = RGBColor(0x59, 0x56, 0x59)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BG_HEADER = "1B3A5C"
COLOR_BG_ROW = "F2F7FB"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def setup_styles(doc):
    """문서 기본 스타일 설정"""
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # 제목 스타일
    for level, (size, color, bold) in {
        "Heading 1": (Pt(22), COLOR_PRIMARY, True),
        "Heading 2": (Pt(16), COLOR_PRIMARY, True),
        "Heading 3": (Pt(13), COLOR_ACCENT, True),
        "Heading 4": (Pt(11), COLOR_GRAY, True),
    }.items():
        s = doc.styles[level]
        s.font.name = FONT_BODY
        s.font.size = size
        s.font.color.rgb = color
        s.font.bold = bold
        s.paragraph_format.space_before = Pt(14 if level == "Heading 1" else 10)
        s.paragraph_format.space_after = Pt(6)
        s.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def add_cover_page(doc):
    """표지 페이지"""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("공존공간")
    run.font.size = Pt(36)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True
    run.font.name = FONT_BODY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("통합 운영 매뉴얼")
    run2.font.size = Pt(28)
    run2.font.color.rgb = COLOR_ACCENT
    run2.font.bold = True
    run2.font.name = FONT_BODY

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("10개 인수인계서 통합본")
    run3.font.size = Pt(14)
    run3.font.color.rgb = COLOR_GRAY
    run3.font.name = FONT_BODY

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("2026. 03. 20.")
    run4.font.size = Pt(12)
    run4.font.color.rgb = COLOR_GRAY
    run4.font.name = FONT_BODY

    doc.add_page_break()


def add_table(doc, headers, rows):
    """테이블 추가"""
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_WHITE
        run.font.name = FONT_BODY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLOR_BG_HEADER)

    # 데이터 행
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx >= len(headers):
                continue
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # bold 처리
            text = val.strip()
            if text.startswith("**") and text.endswith("**"):
                run = p.add_run(text[2:-2])
                run.font.bold = True
            else:
                run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = FONT_BODY
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_BG_ROW)

    doc.add_paragraph()


def parse_table_block(lines, start):
    """마크다운 테이블 파싱. (headers, rows, end_index) 반환"""
    headers = [c.strip() for c in lines[start].split("|")[1:-1]]
    # separator line
    idx = start + 1
    if idx < len(lines) and re.match(r"\|[\s\-:|]+\|", lines[idx]):
        idx += 1
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cols = [c.strip() for c in lines[idx].split("|")[1:-1]]
        rows.append(cols)
        idx += 1
    return headers, rows, idx


def build_doc(md_text):
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    setup_styles(doc)
    add_cover_page(doc)

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # 코드 블록
        if line.strip().startswith("```"):
            if in_code_block:
                # 코드 블록 끝
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.size = Pt(8.5)
                run.font.name = "Consolas"
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 목차, 메타 정보 건너뛰기
        if stripped.startswith("> **통합일"):
            i += 1
            continue
        if stripped.startswith("> **원본"):
            i += 1
            continue

        # 구분선
        if stripped == "---":
            # 챕터 구분 시 페이지 브레이크
            i += 1
            continue

        # 제목
        if stripped.startswith("#"):
            match = re.match(r"^(#{1,4})\s+(.+)", stripped)
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                # H1이면 페이지 브레이크 (첫 번째 제외)
                if level == 1:
                    i += 1
                    continue  # 문서 제목은 표지에서 처리
                if level == 2:
                    doc.add_page_break()
                    doc.add_heading(title, level=1)
                elif level == 3:
                    doc.add_heading(title, level=2)
                elif level == 4:
                    doc.add_heading(title, level=3)
                i += 1
                continue

        # 테이블
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"\|[\s\-:|]+\|", lines[i + 1].strip()):
            headers, rows, end = parse_table_block(lines, i)
            add_table(doc, headers, rows)
            i = end
            continue

        # 목차 링크 건너뛰기
        if re.match(r"^\d+\.\s+\[", stripped):
            i += 1
            continue

        # 블록인용 (> 로 시작)
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            if text.startswith("**참고**"):
                p = doc.add_paragraph()
                run = p.add_run(text.replace("**", ""))
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = COLOR_GRAY
                run.font.name = FONT_BODY
            else:
                text = text.replace("**", "")
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Cm(0.5)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = COLOR_GRAY
            i += 1
            continue

        # 리스트 아이템
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)", stripped)
        if list_match:
            indent = len(list_match.group(1)) if list_match.group(1) else 0
            marker = list_match.group(2)
            text = list_match.group(3)
            is_ordered = bool(re.match(r"\d+\.", marker))

            # bold 처리
            p = doc.add_paragraph(style="List Bullet" if not is_ordered else "List Number")
            p.clear()

            # 간단한 bold 파싱
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(9.5)
                run.font.name = FONT_BODY

            if indent > 2:
                p.paragraph_format.left_indent = Cm(1.5)

            i += 1
            continue

        # 일반 텍스트
        text = stripped
        p = doc.add_paragraph()
        # bold 파싱
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = FONT_BODY

        i += 1

    return doc


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = build_doc(md_text)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
