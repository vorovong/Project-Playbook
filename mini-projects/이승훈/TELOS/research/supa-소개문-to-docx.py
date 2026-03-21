from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = size
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_mixed_para(parts, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        size = part[3] if len(part) > 3 else None
        color = part[4] if len(part) > 4 else None
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if bold: run.bold = True
        if italic: run.italic = True
        if size: run.font.size = size
        if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1A1A2E', qn('w:color'): 'auto', qn('w:val'): 'clear'
        })
        shading.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return table

# =============================================
# v3 최종본
# =============================================

# 타이틀
add_para('슈파(SUPA)', bold=True, size=Pt(24), color=RGBColor(0x1A, 0x1A, 0x2E), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para('사장님이 장사에만 집중할 수 있게, 나머지를 같이 해주는 서비스', size=Pt(12), color=RGBColor(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

add_para('━' * 50, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

# ===== 1막: 공감 =====
add_heading_styled('사장님, 장사 말고 해야 할 일이 너무 많지 않으십니까?', level=2)

add_para('가게 문 열기 전에 식자재 주문.', space_after=Pt(2))
add_para('가게 문 닫고 나서 리뷰 답변, SNS 올리기, 장부 정리.', space_after=Pt(8))

add_mixed_para([
    ('사장님이 진짜 잘하시는 건 ', False),
    ('요리', True),
    ('이고 ', False),
    ('손님 맞이', True),
    ('인데,', False),
])
add_para('정작 그 시간은 줄고 잡무 시간만 늘어갑니다.', space_after=Pt(12))

add_para('혼자가 아닙니다.', bold=True, space_after=Pt(8))

add_para('매출의 절반 가까이가 재료비로 나가고, 남는 건 월 200만 원 남짓이라고 합니다.', space_after=Pt(2))
add_para('문 닫으신 사장님 절반이 "재료비 때문"이라고 하셨습니다.', space_after=Pt(10))

add_para('혼자서 가격 비교하고, 혼자서 주문하고, 혼자서 관리하고 계십니다.', space_after=Pt(12))

add_para('"같이 하면 어떨까?" — 여기서 슈파가 시작됐습니다.', bold=True, size=Pt(12), color=RGBColor(0x1A, 0x1A, 0x2E), space_after=Pt(16))

# ===== 2막: 안도 =====
add_heading_styled('슈파는 사장님 옆에서 같이 일합니다', level=2)

add_para('슈파(SUPA)는 IT 회사가 만든 서비스가 아닙니다.', space_after=Pt(8))

add_mixed_para([
    ('공존공간', True),
    ('이라는 이름으로 지난 2년간 수원에서 사장님들과 ', False),
    ('142번', True),
    (' 직접 만나고, 같이 일하고, 문제를 풀면서 만든 서비스입니다.', False),
], space_after=Pt(8))

add_para('프레젠테이션이 아니라 주방 옆에서 시작했습니다.', bold=True, italic=True, space_after=Pt(12))

# 공존공간 성과 (사장님 언어로 번역)
add_para('그 결과, 함께한 사장님들의 매출이 올랐습니다.', space_after=Pt(4))
add_mixed_para([
    ('행궁동에서 2년간 같이 일한 가게들의 ', False),
    ('고객 소비액이 2.5배', True),
    (', ', False),
    ('사장님 만족도가 64점에서 89점', True),
    ('으로 올랐습니다.', False),
], space_after=Pt(12))

add_para('이 경험을 식자재 구매와 가게 운영으로 확장한 것이 슈파입니다.', space_after=Pt(4))
add_para('사장님이 장사에만 집중할 수 있게, 나머지를 같이 해주는 서비스입니다.', bold=True, space_after=Pt(12))

add_para('그래서 저희는 압니다.', space_after=Pt(2))
add_mixed_para([
    ('사장님한테 필요한 건 "최신 기술"이 아니라 ', False),
    ('"당장 도움 되는 것"', True),
    ('이라는 걸.', False),
], space_after=Pt(16))

# 기능표 — 2단 구분
add_para('지금 쓸 수 있는 것', bold=True, size=Pt(11), color=RGBColor(0x1A, 0x1A, 0x2E), space_after=Pt(6))

add_table(
    ['사장님이 하시던 일', '슈파가 도와드리는 방식'],
    [
        ['식자재 가격 비교, 개별 주문', '공동구매로 같이 사서 단가를 낮춥니다'],
        ['거래명세서 정리, 장부 기록', '엑셀 넣으면 자동 분석해 드립니다'],
    ]
)

doc.add_paragraph()

add_para('곧 함께 만들어갈 것', bold=True, size=Pt(11), color=RGBColor(0x66, 0x66, 0x66), space_after=Pt(6))

add_table(
    ['사장님이 하시던 일', '슈파가 도와드리는 방식'],
    [
        ['리뷰 답변 하나하나 쓰기', '사장님 말투로 답변 초안을 만들어 드립니다'],
        ['SNS 포스팅 고민', '가게 소식 자동 포스팅을 세팅해 드립니다'],
    ]
)

doc.add_paragraph()

# ===== 3막: 확신 =====
add_heading_styled('사장님 가게가 이렇게 달라집니다', level=2)

add_para('사장님께서 혹시 이런 경험 있으시지 않습니까.', space_after=Pt(4))
add_para('"무료"라고 해서 써봤더니, 나중에 수수료 붙고, 결국 또 돈 나가는.', italic=True, space_after=Pt(10))

add_para('그래서 먼저 결과를 보여드리겠습니다.', space_after=Pt(8))
add_para('가게 하나가 달라지는 모습을 그려보겠습니다.', bold=True, size=Pt(12), color=RGBColor(0x1A, 0x1A, 0x2E), space_after=Pt(10))

# 시나리오 — 출처 괄호 없이 서술체
add_para('아까 그 200만 원, 어디서 나온 숫자인지 한번 보겠습니다.', space_after=Pt(6))
add_para('월 매출 2,700만 원짜리 가게가 있습니다.', space_after=Pt(2))
add_para('이 중 절반 가까이, 약 1,100만 원이 식재료비입니다.', space_after=Pt(2))
add_mixed_para([
    ('공동구매로 ', False),
    ('5%만 줄이면 월 55만 원, 연 660만 원', True),
    ('이 남습니다.', False),
], space_after=Pt(2))
add_mixed_para([
    ('이 가게 순이익이 200만 원대라면, 55만 원은 ', False),
    ('순이익 25% 증가', True),
    ('입니다.', False),
], space_after=Pt(12))

add_para('리뷰 답변에 매일 30분 쓰셨다면, 슈파가 사장님 말투로 초안을 만들어 드립니다.', space_after=Pt(2))
add_mixed_para([
    ('확인하고 전송만 누르시면 됩니다. ', False),
    ('그 30분을 손님한테 쓰실 수 있습니다.', True),
], space_after=Pt(12))

add_para('지금 비용을 받지 않는 이유는 분명합니다.', space_after=Pt(2))
add_mixed_para([
    ('저희가 수원에서 먼저 검증하는 단계이고, ', False),
    ('사장님들과 같이 만들어가야 하는 서비스', True),
    ('이기 때문입니다.', False),
], space_after=Pt(16))

# ===== 4막: 행동 =====
add_heading_styled('지금 시작하는 30곳이 가장 큰 혜택을 받습니다', level=2)

add_mixed_para([
    ('슈파는 ', False),
    ('수원 지역 시범 파트너 30곳', True),
    ('을 모집합니다.', False),
], space_after=Pt(12))

add_para('시범 파트너에게 드리는 것:', bold=True, space_after=Pt(6))

benefits = [
    ('1. 시범 기간 완전 무료', ' — 비용 부담 없이 써보실 수 있습니다'),
    ('2. 가게 맞춤 세팅', ' — 사장님 가게 상황을 듣고, 필요한 것부터 세팅해 드립니다'),
    ('3. 우선 반영', ' — 사장님의 의견이 서비스에 바로 반영됩니다'),
]
for b_bold, b_normal in benefits:
    add_mixed_para([
        (b_bold, True),
        (b_normal, False),
    ], space_after=Pt(3))

doc.add_paragraph()

add_para('시작은 이렇게 간단합니다', bold=True, size=Pt(12), space_after=Pt(10))

# 카카오톡 집중
add_para("카카오톡에서 '슈파'를 검색하시고, '상담'이라고 한 마디만 보내주세요.", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10))

add_para('사장님 가게에 직접 찾아갑니다. 30분이면 됩니다 — 영업 중에도 괜찮습니다.', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para('그 자리에서 바로 세팅해 드립니다.', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))

add_para('(전화를 선호하시면 000-0000-0000으로 편하게 연락 주십시오)', size=Pt(9), italic=True, color=RGBColor(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

# 구분선
add_para('━' * 50, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

# 마무리 멘트
add_para('가게 이야기, 편하게 들려주세요.', bold=True, size=Pt(14), color=RGBColor(0x1A, 0x1A, 0x2E), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para('거기서부터 시작하겠습니다.', bold=True, size=Pt(14), color=RGBColor(0x1A, 0x1A, 0x2E), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

# 각주 — 출처 모음
add_para('━' * 50, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para('참고 수치 출처', bold=True, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99), space_after=Pt(3))
sources = [
    '외식업 식재료비 비중 40.4% — 농림축산식품부, 2024 외식업체 경영실태조사',
    '외식업 월 영업이익 221만원 — 식품외식경제, 2024',
    '폐업 원인 중 원재료비 부담 46% — 중소기업중앙회, 2025 폐업 소상공인 실태조사',
    '외식 가맹점 평균 연매출 3억 2,300만원 — 푸드뉴스, 2024',
    '공존공간 소상공인 142회 협업·고객 소비액·만족도 — 공존공간 성과보고, 2023~2025',
]
for s in sources:
    add_para(s, size=Pt(7), color=RGBColor(0xAA, 0xAA, 0xAA), space_after=Pt(1))

# 저장
output_path = r'C:\Users\leeha\projects\Project_Playbook\mini-projects\이승훈\TELOS\research\슈파-소개문-v3-최종.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
