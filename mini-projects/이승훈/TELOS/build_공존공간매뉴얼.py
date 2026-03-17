# -*- coding: utf-8 -*-
"""공존공간 운영매뉴얼 Word 문서 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 기본 스타일 설정 ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        # 헤더 배경색
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2C3E50',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            # 짝수행 배경
            if r_idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F3F5',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elm)

    doc.add_paragraph()  # 간격
    return table

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'※ {text}')
    run.font.name = '맑은 고딕'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    run.italic = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


# ════════════════════════════════════════
# 문서 본문 시작
# ════════════════════════════════════════

add_title('공존공간 운영 매뉴얼')
doc.add_paragraph()

# ── 1. 경영지원 ──
add_h1('1. 경영지원')

# 1-1. 인사 업무
add_h2('1-1. 인사 업무')
add_table(
    ['업무', '상세'],
    [
        ['구인 공고 등록', '잡코리아·사람인 활용\n기존 공고 기준으로 변경사항만 반영하여 작성'],
        ['지원자 관리', '이력서·포트폴리오를 PDF 저장 후 보고\n신규 채용 시 루틴업무로 전환 → 매일 오전 구인구직 톡방에 보고'],
        ['면접 일정 안내', '일정 확정 후 안내 문자 발송\n010-4188-2117 핸드폰 기록 확인 → 변경 내용만 수정하여 발송'],
        ['입사 전 안내', '입사 일정 안내 + 입사 전 참고자료 메일 송부'],
        ['입사 후 인계', '입퇴사자 안내 탭 내용에 따라 인계 진행'],
    ]
)

# 1-2. 근태 및 월차 관리
add_h2('1-2. 근태 및 월차 관리')
add_table(
    ['업무', '상세'],
    [
        ['월차 관리', '월차대장 파일에서 직원별 관리\n1개월 만근 시 월차 1개 발생'],
        ['대체휴무', '발생 시 → "대휴"로 기입\n사용 시 → "대발"로 기입'],
        ['월차 사용 보고', '부득이한 사유 제외, 사용 2주 전 대표님 보고 필요\n(수빈님 진행)'],
        ['근태 조회', '캡스 로그인 → 상단 근태 탭 → 조회\n→ 조건 선택 후 엑셀 or PDF 다운'],
    ]
)

# 1-3. 급여 및 4대보험
add_h2('1-3. 급여 및 4대보험')
add_note('내용 추가 필요')

# 1-4. 법정 의무 교육
add_h2('1-4. 법정 의무 교육')
add_note('현재 진행 여부 확인 필요')

# 1-5. 국비·지방비 사업비 관리
add_h2('1-5. 국비·지방비 사업비 관리')
add_bullet('수원시 지방비')
add_note('종료 여부 확인 필요')

# 1-6. 월 정산 업무
add_h2('1-6. 월 정산 업무')
add_table(
    ['업무', '상세'],
    [
        ['공유오피스', '세금계산서 발행'],
        ['미식가의주방', '매주 정산, 매월 정산, 공과금 관리\n→ 미식가 대표님과 소통'],
        ['세무사무실', '요청 자료 주고받기\n(법인세, 원천세 등)'],
        ['양조장', '세금계산서 발행 및 입금 확인'],
    ]
)

# ── 2. 공간 운영 ──
add_h1('2. 공간 운영')

add_h2('2-1. 양조장 관리')
add_note('상세 내용 추가 필요')

add_h2('2-2. 재생전술 원데이 클래스')
add_bullet('예약금 입금 확인 및 환불 처리')

add_h2('2-3. 포토인더박스')
add_bullet('환불 처리')

add_h2('2-4. 사옥 청결화 관리')
add_note('상세 내용 추가 필요')

# ── 3. 입퇴사 체크리스트 ──
add_h1('3. 입퇴사 체크리스트')

# 3-1. 입사 절차
add_h2('3-1. 입사 절차')

add_h3('① 시스템 등록')
add_table(
    ['순서', '업무', '방법'],
    [
        ['1', '채널톡 가입', '채널톡 매니저 초대 링크 전송 → 가입 요청'],
        ['2', '구글 계정 배정', '보안용(2025 웹사이트 id/pw)\n+ 실무용(2025 공존 웹사이트 id/pw)\n내부 확인 후 각각 전달'],
        ['3', '나스 등록', '—'],
        ['4', '지문 등록', '캡스 프로그램 실행 → 지문등록관리\n→ 마스터관리자 로그온 (비밀번호: caps!)\n→ 등록관리 → 사용자 관리\n→ 이름 작성 → 지문 등록(최대 2개)\n→ 추가 → 사용자 목록 확인'],
    ]
)

add_h3('② 계약 진행')
add_table(
    ['순서', '업무', '방법'],
    [
        ['1', '간이 계약', '입사일 기준 1주일간 간이 계약 진행\n⚠ 반드시 사전 고지 필수!\n통장 사본·신분증 사본 요청'],
        ['2', '근로계약서', '표준근로계약서 양식으로 작성 → 2부 인쇄\n→ 서명 → 원본 2부 전체 인쇄 후 저장\n→ 1부 보관, 1부 당사자 전달'],
    ]
)

# 3-2. 퇴사 절차
add_h2('3-2. 퇴사 절차')
add_table(
    ['순서', '업무'],
    [
        ['1', '사직서 작성'],
        ['2', '원천징수영수증 전달'],
        ['3', '이직확인서 등록'],
    ]
)

# ── 저장 ──
import subprocess
result = subprocess.run(
    ['powershell', '-Command', '[Environment]::GetFolderPath("Desktop")'],
    capture_output=True, text=True
)
desktop = result.stdout.strip()
output_path = os.path.join(desktop, 'gongjon_manual.docx')
doc.save(output_path)
print(f'저장 완료: {output_path}')
